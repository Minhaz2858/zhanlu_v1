"""Verify the artifact GET endpoint advertises DOCX preview modes + outline."""
import io

from docx import Document
from fastapi.testclient import TestClient

from main import app
from app.database import SessionLocal
from app.services.artifacts.artifact_service import ArtifactService


def _seed_docx_artifact(
    db, headings=("Executive Summary", "Method"), name="test.docx"
) -> str:
    svc = ArtifactService(db)
    art = svc.create_artifact(artifact_type="docx", title="Test Plan")
    ver = svc.create_version(artifact_id=art.id)
    doc = Document()
    doc.add_heading(headings[0], level=1)
    doc.add_paragraph("Body")
    if len(headings) > 1:
        doc.add_heading(headings[1], level=2)
        doc.add_paragraph("More body")
    buf = io.BytesIO()
    doc.save(buf)
    svc.store_blob(
        version_id=ver.id,
        blob_type="original",
        file_name=name,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buf.getvalue(),
    )
    return art.id


def test_artifact_payload_includes_preview_modes_and_outline(monkeypatch):
    from app.config import settings as app_settings
    original = app_settings.APP_PUBLIC_URL
    monkeypatch.setattr(app_settings, "APP_PUBLIC_URL", "")  # no public URL → only self_hosted
    try:
        with SessionLocal() as db:
            artifact_id = _seed_docx_artifact(db)
        client = TestClient(app)
        resp = client.get(f"/api/artifacts/{artifact_id}")
        assert resp.status_code == 200, resp.text
        body = resp.json()
        assert body["artifact_type"] == "docx"
        assert body["preview_modes"] == ["self_hosted_html"]
        assert body["preview_outline"] is not None
        # Check outline structure
        outline = body["preview_outline"]
        assert isinstance(outline, list)
        assert len(outline) >= 1
        assert all("level" in o and "text" in o and "id" in o for o in outline)
    finally:
        app_settings.APP_PUBLIC_URL = original


def test_artifact_payload_ms_word_open_url_when_public_url_set(monkeypatch):
    """With APP_PUBLIC_URL set, ms_word_open_url should be present and non-null."""
    from app.config import settings as app_settings
    original = app_settings.APP_PUBLIC_URL
    monkeypatch.setattr(app_settings, "APP_PUBLIC_URL", "https://zhanlu.example.com")
    try:
        with SessionLocal() as db:
            artifact_id = _seed_docx_artifact(db, name="plan.docx")
        client = TestClient(app)
        resp = client.get(f"/api/artifacts/{artifact_id}")
        assert resp.status_code == 200, resp.text
        body = resp.json()
        assert "ms_word" in body["preview_modes"]
        assert "self_hosted_html" in body["preview_modes"]
        assert body["ms_word_open_url"] is not None
        assert body["ms_word_open_url"].startswith(
            "https://view.officeapps.live.com/op/embed.aspx?src="
        )
    finally:
        app_settings.APP_PUBLIC_URL = original


def test_non_docx_has_empty_preview_fields():
    """Non-DOCX artifacts should have empty preview_modes and outline."""
    with SessionLocal() as db:
        svc = ArtifactService(db)
        art = svc.create_artifact(artifact_type="md", title="README")
        svc.create_version(artifact_id=art.id)
        artifact_id = art.id
    client = TestClient(app)
    resp = client.get(f"/api/artifacts/{artifact_id}")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["preview_modes"] == []
    assert body["preview_outline"] == []
    assert body.get("ms_word_open_url") is None
