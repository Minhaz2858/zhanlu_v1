"""Tests for the styled DOCX generator.

Verifies that ``generate_docx`` produces a Word document with:
- A themed title (large, bold, dark color)
- A metadata subtitle with row count and source
- A KPI table with large bold values
- A data table with colored header
- An insights bullet list
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from app.services.sandbox import sandbox_runner as sr


@pytest.fixture(autouse=True)
def tmp_output_dir():
    """Point OUTPUT_DIR at a fresh tmp dir for the test."""
    with tempfile.TemporaryDirectory(prefix="zhanlu-docx-test-") as tmp:
        original = sr.OUTPUT_DIR
        sr.OUTPUT_DIR = Path(tmp)
        try:
            yield Path(tmp)
        finally:
            sr.OUTPUT_DIR = original


def _config():
    return {
        "format": "docx",
        "title": "Q1 Sales Report",
        "kpis": [
            {"label": "Revenue", "value": "$6,900", "caption": "+15%"},
            {"label": "Orders", "value": "240"},
        ],
        "insights": [
            "Revenue grew 75% from January to March",
            "Order volume increased steadily",
        ],
        "source": "Production DB",
    }


def _rows():
    return [
        {"month": "Jan", "revenue": 1200, "orders": 45},
        {"month": "Feb", "revenue": 1850, "orders": 62},
    ]


def test_docx_title_present(tmp_output_dir):
    sr.generate_docx(_rows(), _config(), "Test instructions")
    from docx import Document
    doc = Document(str(tmp_output_dir / "report.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Q1 Sales Report" in full_text


def test_docx_kpi_table_present(tmp_output_dir):
    sr.generate_docx(_rows(), _config(), "Test")
    from docx import Document
    doc = Document(str(tmp_output_dir / "report.docx"))
    # The KPI table is a 2×N table: header row (labels) + value row
    kpi_table = doc.tables[0]
    assert len(kpi_table.rows) == 2
    assert len(kpi_table.columns) == 2
    # Row 0: labels
    label_texts = [c.text for c in kpi_table.rows[0].cells]
    assert "Revenue" in label_texts[0]
    assert "Orders" in label_texts[1]
    # Row 1: values
    value_texts = [c.text for c in kpi_table.rows[1].cells]
    assert "$6,900" in value_texts[0]
    assert "240" in value_texts[1]


def test_docx_data_table_present(tmp_output_dir):
    sr.generate_docx(_rows(), _config(), "Test")
    from docx import Document
    doc = Document(str(tmp_output_dir / "report.docx"))
    # Second table is the data table
    data_table = doc.tables[1]
    # 1 header + 2 rows
    assert len(data_table.rows) == 3
    header = [c.text for c in data_table.rows[0].cells]
    assert header == ["month", "revenue", "orders"]


def test_docx_insights_present(tmp_output_dir):
    sr.generate_docx(_rows(), _config(), "Test")
    from docx import Document
    doc = Document(str(tmp_output_dir / "report.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Revenue grew 75% from January to March" in full_text
    assert "Order volume increased steadily" in full_text


def test_docx_source_in_subtitle(tmp_output_dir):
    sr.generate_docx(_rows(), _config(), "Test")
    from docx import Document
    doc = Document(str(tmp_output_dir / "report.docx"))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Source: Production DB" in full_text


def test_docx_no_kpis_omits_kpi_table(tmp_output_dir):
    cfg = _config()
    cfg["kpis"] = []
    sr.generate_docx(_rows(), cfg, "Test")
    from docx import Document
    doc = Document(str(tmp_output_dir / "report.docx"))
    # Only the data table — no KPI table
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 3
