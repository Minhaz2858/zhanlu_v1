"""Automation DOCX quality: the 3-tier renderer produces native Word tables
and inline formatting, and the report prompt forbids technical metadata."""
import io
import os
import sys
from pathlib import Path

_BACKEND_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _BACKEND_ROOT not in sys.path:
    sys.path.insert(0, _BACKEND_ROOT)
os.chdir(_BACKEND_ROOT)


MARKDOWN_SAMPLE = """## Executive summary

本次增量同步未发现新增订单。

## Key metrics

| 指标 | 数值 | 较上次运行 |
|---|---|---|
| 增量订单数 | 0 单 | 持平 |
| 增量金额 | ¥0.00 | 持平 |

## Changes since last run

- **No changes detected**

## Recommended actions

1. 排查上游同步链路
2. 复核增量过滤口径
"""


def _read_docx_bytes(docx_bytes: bytes):
    from docx import Document
    return Document(io.BytesIO(docx_bytes))


def test_via_python_docx_preserves_tables_and_bold():
    from app.services.artifacts.exporters.html_docx import render_html_to_docx

    html = (
        "<html><body><h1>Report</h1>"
        "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td><strong>bold</strong></td></tr></table>"
        "<p>This is <strong>bold</strong> and <em>italic</em> text.</p>"
        "</body></html>"
    )
    docx_bytes = render_html_to_docx(html.encode("utf-8"))
    assert docx_bytes and len(docx_bytes) > 0

    doc = _read_docx_bytes(docx_bytes)
    # Native Word table, not pipe text.
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "A"
    assert table.rows[0].cells[1].text == "B"
    assert table.rows[1].cells[1].text == "bold"

    # Inline bold/italic are preserved as runs (no literal asterisks).
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "**" not in all_text
    assert "bold" in all_text and "italic" in all_text
    assert any(r.bold for p in doc.paragraphs for r in p.runs)


def test_generate_document_docx_produces_native_table(tmp_path, monkeypatch):
    from app.services import document_generator as dg

    # Redirect output to the test's temp dir — the real data/generated tree is
    # owned by the container uid and not writable from the test user.
    root = tmp_path / "automation"
    root.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(dg, "_automation_root", lambda: root)

    path, url, mime = dg.generate_document(
        output_format="docx",
        content=MARKDOWN_SAMPLE,
        title="ERP Sales Sync",
        task_id="t1",
        exec_id="e1",
    )
    assert path.exists()
    assert mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    doc = _read_docx_bytes(path.read_bytes())
    assert len(doc.tables) >= 1
    # Header cells present, not raw pipe text.
    header_text = " ".join(c.text for c in doc.tables[0].rows[0].cells)
    assert "指标" in header_text
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "|" not in full_text


def test_document_service_parser_is_exported():
    from app.services.document_service import parse_markdown_to_docx
    from docx import Document

    doc = Document()
    parse_markdown_to_docx(doc, MARKDOWN_SAMPLE)
    # Should produce a table for the key metrics section.
    assert len(doc.tables) >= 1


def test_report_prompt_forbids_technical_metadata():
    from app.services.automation_executor import _REPORT_STRUCTURE_GUIDANCE

    g = _REPORT_STRUCTURE_GUIDANCE.lower()
    for banned in (
        "execution ids",
        "database hostnames",
        "connection strings",
        "retry policies",
        "raw sql",
        "business-facing",
    ):
        assert banned in g
