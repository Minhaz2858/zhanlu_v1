"""GET /api/artifacts/{id}/preview?format=html returns sanitized DOCX HTML."""
import io

from docx import Document
from fastapi.testclient import TestClient

from main import app
from app.database import SessionLocal
from app.services.artifacts.artifact_service import ArtifactService


def _seed() -> str:
    svc = ArtifactService(SessionLocal())
    art = svc.create_artifact(artifact_type="docx", title="Hello")
    ver = svc.create_version(artifact_id=art.id)
    doc = Document()
    doc.add_heading("Hi", level=1)
    doc.add_paragraph("Line <b>bold</b>")
    buf = io.BytesIO()
    doc.save(buf)
    svc.store_blob(
        version_id=ver.id,
        blob_type="original",
        file_name="hello.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buf.getvalue(),
    )
    return art.id


def test_preview_html_returns_sanitized_html():
    artifact_id = _seed()
    client = TestClient(app)
    resp = client.get(f"/api/artifacts/{artifact_id}/preview?format=html")
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("text/html")
    body = resp.text
    assert "<h1" in body
    # raw HTML inside the doc must be escaped
    assert "&lt;b&gt;bold&lt;/b&gt;" in body
    assert "<script>" not in body


def test_preview_html_400_for_unsupported_format():
    artifact_id = _seed()
    client = TestClient(app)
    resp = client.get(f"/api/artifacts/{artifact_id}/preview?format=garbage")
    assert resp.status_code in (400, 404)


def test_preview_html_400_for_non_docx_html():
    """?format=html on a non-DOCX artifact should return 400."""
    svc = ArtifactService(SessionLocal())
    art = svc.create_artifact(artifact_type="md", title="README")
    svc.create_version(artifact_id=art.id)
    client = TestClient(app)
    resp = client.get(f"/api/artifacts/{art.id}/preview?format=html")
    assert resp.status_code == 400, resp.text
    assert "html" in resp.json()["detail"].lower()
