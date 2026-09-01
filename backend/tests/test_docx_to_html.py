"""Unit tests for convert_docx_to_html and extract_docx_outline."""
import io

import pytest
from docx import Document

from app.services.artifacts.preview_builder import convert_docx_to_html, extract_docx_outline


def _make_docx(headings=("Title", "Section A", "Section B"), body=("Body paragraph.",)):
    doc = Document()
    doc.add_heading(headings[0], level=0)
    for h in headings[1:]:
        doc.add_heading(h, level=1)
    for line in body:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_convert_docx_to_html_strips_xml():
    html, messages = convert_docx_to_html(_make_docx())
    assert "<h1" in html
    assert "Body paragraph." in html
    # mammoth messages are warnings; we just want a list back
    assert isinstance(messages, list)


def test_convert_docx_to_html_escapes_raw_html_in_text():
    """Plain paragraphs must be HTML-escaped (no XSS in inline preview)."""
    raw = _make_docx(body=("<script>alert(1)</script>",))
    html, _ = convert_docx_to_html(raw)
    assert "<script>" not in html
    assert "&lt;script&gt;" in html


def test_convert_docx_to_html_returns_empty_on_bad_bytes():
    html, messages = convert_docx_to_html(b"not a real docx")
    assert html == ""
    assert any("mammoth" in (m or "").lower() or "error" in (m or "").lower()
               for m in messages)


def test_extract_docx_outline_returns_headings_in_order():
    outline = extract_docx_outline(_make_docx(headings=("Title", "Alpha", "Beta", "Gamma")))
    texts = [o["text"] for o in outline]
    assert texts == ["Title", "Alpha", "Beta", "Gamma"]


def test_extract_docx_outline_handles_empty_doc():
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    assert extract_docx_outline(buf.getvalue()) == []
