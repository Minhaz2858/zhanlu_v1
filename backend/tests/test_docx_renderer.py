"""Tests for the enterprise DOCX renderer (spec §10.1)."""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.enterprise_orchestrator.renderers.docx import (
    render_enterprise_docx,
)


def _sample_payload() -> dict:
    return {
        "enterprise_report_kind": "executive",
        "title": "2026-07-21 to 2026-08-20 Supply Chain Report",
        "domain": "supply_chain",
        "period": ["2026-07-21", "2026-08-20"],
        "period_label": "2026-07-21 to 2026-08-20",
        "primary_metric": "volume",
        "executive_summary": (
            "Supply Chain summary for 2026-07-21 to 2026-08-20: "
            "1,234.56 tons sold; top-3 customer share 45.0%."
        ),
        "primary_metric_breakdown": {
            "label": "Primary Metric Breakdown",
            "metric": "volume",
            "rows": [
                {"material_name": "Steel", "total_volume_tons": 500.0, "total_revenue": 50.0},
                {"material_name": "Aluminium", "total_volume_tons": 400.0, "total_revenue": 40.0},
            ],
            "kpi": {"total_volume_tons": 1234.56, "total_revenue": 123.0},
            "available": True,
            "unavailable_reason": "",
        },
        "segment_decomposition": {
            "label": "Segment Decomposition",
            "rows": [
                {"customer_name": "Acme", "total_volume_tons": 300.0, "total_revenue": 30.0},
                {"customer_name": "Globex", "total_volume_tons": 200.0, "total_revenue": 20.0},
            ],
            "available": True,
            "unavailable_reason": "",
            "observations": ["Top segment: Acme with 300.00 tons / ¥30.00M revenue."],
        },
        "operational_drivers": {
            "label": "Supply-Demand Drivers",
            "available": True,
            "drivers": ["Volume + revenue dominated by top 5 materials."],
        },
        "risk_section": {
            "label": "Transmission Risk",
            "available": True,
            "risks": ["Customer concentration high: top-3 = 45.00%."],
        },
        "recommended_actions": [
            {
                "action": "Restock SKUs with sub-five-day supply before the next delivery window.",
                "severity": "high",
                "source_facet": "inventory_position",
                "metric_value": 2.0,
                "threshold": 5.0,
            },
        ],
        "data_confidence": {
            "covered_facets": ["sales_summary", "inventory_position"],
            "missing_facets": ["competitor_benchmark"],
            "missing_reasons": {"competitor_benchmark": "no competitor data table"},
        },
        "claims": [
            {
                "claim_id": "primary_kpi",
                "text": "Total volume over the period: 1,234.56 tons",
                "source_facet": "sales_summary",
                "source_row_ids": ["Steel", "Aluminium"],
                "source_sql": "SELECT SUM(qty) FROM erp_v_sale_orderentry",
                "verified": True,
            },
        ],
        "lineage": {
            "sales_summary": {
                "source_sql": "SELECT SUM(qty) FROM erp_v_sale_orderentry",
                "source_label": "erp_v_sale_orderentry",
                "row_count": 2,
                "execution_log": [
                    {"step": "service_invoke", "latency_ms": 120, "status": "ok"},
                ],
                "available": True,
            },
            "competitor_benchmark": {
                "source_sql": "",
                "source_label": "",
                "row_count": 0,
                "execution_log": [],
                "available": False,
            },
        },
        "created_at": "2026-08-20T00:00:00Z",
    }


def _render(payload: dict | None = None) -> str:
    doc = Document(io.BytesIO(render_enterprise_docx(payload or _sample_payload())))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def test_returns_bytes():
    out = render_enterprise_docx(_sample_payload())
    assert isinstance(out, bytes)
    assert len(out) > 0


def test_all_six_sections_present():
    text = _render()
    for heading in (
        "Executive Summary",
        "Primary Metric Breakdown",
        "Segment Decomposition",
        "Supply-Demand Drivers",
        "Transmission Risk",
        "Recommended Actions",
    ):
        assert heading in text, f"missing section heading: {heading}"


def test_cover_page_has_title_and_period():
    text = _render()
    assert "2026-07-21 to 2026-08-20 Supply Chain Report" in text
    assert "2026-07-21 to 2026-08-20" in text


def test_title_is_not_generic_executive_summary():
    # The document title must be the domain/period title, never the
    # generic "Executive Summary" heading (which is only a section H1).
    payload = _sample_payload()
    doc = Document(io.BytesIO(render_enterprise_docx(payload)))
    # The first non-empty paragraph is the title on the cover page.
    first_text = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    assert first_text == payload["title"]
    assert "Executive Summary" not in first_text


def test_lineage_appendix_has_sql_and_execution_log():
    text = _render()
    assert "Data Lineage" in text
    assert "SELECT SUM(qty) FROM erp_v_sale_orderentry" in text
    assert "service_invoke" in text
    assert "erp_v_sale_orderentry" in text


def test_missing_facet_rendered_as_unavailable():
    text = _render()
    assert "competitor_benchmark" in text
    assert "no competitor data table" in text


def test_primary_metric_table_renders_rows():
    text = _render()
    assert "Steel" in text
    assert "Aluminium" in text


def test_recommended_actions_rendered():
    text = _render()
    assert "Restock SKUs" in text
    assert "high" in text


def test_claim_citation_superscript_present():
    # Executive summary citations render as bracketed markers [1].
    text = _render()
    assert "[1]" in text
