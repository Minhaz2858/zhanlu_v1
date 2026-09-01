"""Tests for the Claude-style DOCX export.

These tests exercise the new fields added to ``ReportCardPayload``
(``methodology``, ``key_findings``, ``recommendations``, ``sections``,
``sql``) and verify that ``docx_export.render`` produces a docx
with all the Claude-style sections (cover page, executive summary,
methodology, KPI table, key findings, insights, recommendations,
custom sections, SQL, next step, page-number footer).
"""
import io

import pytest

from app.services.artifacts.exporters.docx_export import (
    render,
)
from app.services.synexia.contracts import (
    ChartSpec,
    InsightSpec,
    KPISpec,
    ReportCardPayload,
    SectionSpec,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _full_payload(**overrides) -> ReportCardPayload:
    """Build a ReportCardPayload with all Claude-style fields populated."""
    defaults = dict(
        title="Q3 Sales Performance Review",
        source="sales_db",
        generated_at="2026-07-15T10:00:00",
        summary="Revenue grew 23% YoY driven by Q3 product launches and EMEA expansion.",
        methodology=(
            "Aggregated from sales_db.transactions, deduplicated on customer_id, "
            "filtered to the last 90 days, with currency normalization to USD."
        ),
        kpis=[
            KPISpec(label="Total Revenue", value="$1.2M", delta="+23%", caption="YoY"),
            KPISpec(label="Active Customers", value="1,847", delta="+8%"),
            KPISpec(label="Avg Order", value="$649", caption="+15% YoY"),
        ],
        chart=ChartSpec(
            title="Revenue by Quarter",
            type="bar",
            x_key="quarter",
            y_keys=["revenue"],
            data=[
                {"quarter": "Q1", "revenue": 280_000},
                {"quarter": "Q2", "revenue": 320_000},
                {"quarter": "Q3", "revenue": 380_000},
            ],
        ),
        key_findings=[
            InsightSpec(text="Q3 product launch drove a 19% MoM increase in conversion.", icon="trending-up"),
            InsightSpec(text="EMEA remains the largest market with 44% of total revenue.", icon="globe"),
            InsightSpec(text="APAC grew 34% QoQ — the fastest of any region.", icon="trending-up"),
        ],
        insights=[
            InsightSpec(text="APAC grew 34% QoQ.", icon="trending-up"),
        ],
        recommendations=[
            InsightSpec(text="Increase Q4 inventory by 20% to meet projected demand.", icon="target"),
            InsightSpec(text="Hire 2 more AEs in EMEA to capitalize on expansion momentum.", icon="users"),
        ],
        sections=[
            SectionSpec(title="Data Source", content="All metrics pulled from sales_db.daily_transactions."),
            SectionSpec(
                title="Caveats",
                bullets=["Returns not netted out", "Q4 forecast assumes no macro shocks"],
            ),
        ],
        sql="SELECT quarter, SUM(amount) AS revenue FROM sales WHERE year=2026 GROUP BY quarter;",
        next_step="Plan Q4 inventory based on the Q3 demand pattern.",
        user_signal="export",
    )
    defaults.update(overrides)
    return ReportCardPayload(**defaults)


def _docx_bytes(payload: ReportCardPayload) -> bytes:
    """Render the payload to docx bytes and return them."""
    data, mime, ext = render(payload)
    return data


def _docx_doc(data: bytes):
    """Open a docx bytes payload as a python-docx Document."""
    from docx import Document
    return Document(io.BytesIO(data))


# ---------------------------------------------------------------------------
# Render-level tests
# ---------------------------------------------------------------------------


def test_full_payload_produces_a_non_empty_docx():
    """The happy path: a full payload renders a real docx file."""
    data = _docx_bytes(_full_payload())
    assert len(data) > 20_000  # a Claude-style docx is sizeable
    # The file should start with the ZIP magic (PK..) — docx is a zip.
    assert data[:2] == b"PK"


def test_render_returns_application_vnd_mime():
    from app.services.artifacts.exporters.docx_export import MIME, EXT
    assert MIME == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert EXT == ".docx"


def test_cover_page_contains_title_and_source():
    """The first page should display the report title, source, and date."""
    data = _docx_bytes(_full_payload(title="MY BIG REPORT", source="orders_db"))
    doc = _docx_doc(data)
    # Title + source + date + brand are all on the cover.
    all_paragraphs = "\n".join(p.text for p in doc.paragraphs)
    assert "MY BIG REPORT" in all_paragraphs
    assert "orders_db" in all_paragraphs
    # The title should be centered & large — easiest proxy: the very
    # first non-empty paragraph contains the emoji marker.
    first = next(p for p in doc.paragraphs if p.text.strip())
    assert "📊" in first.text


def test_executive_summary_section_is_present():
    data = _docx_bytes(_full_payload(summary="This is the executive summary line."))
    doc = _docx_doc(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Executive Summary" in text
    assert "This is the executive summary line." in text


def test_methodology_section_is_present():
    data = _docx_bytes(_full_payload(methodology="Query joined 4 tables, deduped, normalized."))
    doc = _docx_doc(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Methodology" in text
    assert "joined 4 tables" in text


def test_key_findings_rendered_as_paragraphs_not_bullets():
    data = _docx_bytes(_full_payload(key_findings=[
        InsightSpec(text="Finding one paragraph."),
        InsightSpec(text="Finding two paragraph."),
    ]))
    doc = _docx_doc(data)
    paragraphs = doc.paragraphs
    # Each finding should be a Normal paragraph, NOT a List Bullet.
    for f in ("Finding one paragraph.", "Finding two paragraph."):
        matches = [p for p in paragraphs if p.text == f]
        assert len(matches) == 1
        assert matches[0].style.name == "Normal", \
            f"Key findings should be paragraphs, got {matches[0].style.name}"


def test_insights_rendered_as_bullets():
    data = _docx_bytes(_full_payload(insights=[
        InsightSpec(text="Insight A"),
        InsightSpec(text="Insight B"),
    ]))
    doc = _docx_doc(data)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    bullet_texts = {b.text for b in bullets}
    assert "Insight A" in bullet_texts
    assert "Insight B" in bullet_texts


def test_recommendations_rendered_as_bullets():
    data = _docx_bytes(_full_payload(recommendations=[
        InsightSpec(text="Hire two more AEs"),
        InsightSpec(text="Expand to APAC"),
    ]))
    doc = _docx_doc(data)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    bullet_texts = {b.text for b in bullets}
    assert "Hire two more AEs" in bullet_texts
    assert "Expand to APAC" in bullet_texts


def test_custom_sections_render_in_declared_order():
    """Sections should appear in the order they were given."""
    data = _docx_bytes(_full_payload(sections=[
        SectionSpec(title="Zebra Section", content="Z content"),
        SectionSpec(title="Apple Section", content="A content"),
    ]))
    doc = _docx_doc(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Zebra should appear before Apple
    z_idx = text.index("Zebra Section")
    a_idx = text.index("Apple Section")
    assert z_idx < a_idx


def test_custom_section_with_bullets():
    data = _docx_bytes(_full_payload(sections=[
        SectionSpec(title="Caveats", bullets=["Returns not netted", "Forecast assumes no shocks"]),
    ]))
    doc = _docx_doc(data)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    bullet_texts = {b.text for b in bullets}
    assert "Returns not netted" in bullet_texts
    assert "Forecast assumes no shocks" in bullet_texts


def test_sql_block_renders_in_a_heading():
    data = _docx_bytes(_full_payload(sql="SELECT 1 FROM dual"))
    doc = _docx_doc(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "SQL" in text
    assert "SELECT 1 FROM dual" in text


def test_next_step_is_not_rendered_in_report():
    # next_step is conversational guidance for the in-chat card, not report
    # content — it must never appear in the exported .docx.
    data = _docx_bytes(_full_payload(next_step="Re-run the analysis weekly."))
    doc = _docx_doc(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Next Step" not in text
    assert "Re-run the analysis weekly." not in text


def test_kpi_table_has_label_and_value_rows():
    data = _docx_bytes(_full_payload(kpis=[
        KPISpec(label="Revenue", value="$1M"),
        KPISpec(label="Customers", value="1,000"),
    ]))
    doc = _docx_doc(data)
    # Find a 2-row, 2-col table (header + values, 2 KPIs)
    kpi_table = next(
        t for t in doc.tables
        if len(t.rows) == 2 and len(t.columns) == 2
    )
    labels = [kpi_table.rows[0].cells[c].text for c in range(2)]
    values = [kpi_table.rows[1].cells[c].text for c in range(2)]
    assert labels == ["Revenue", "Customers"]
    assert values == ["$1M", "1,000"]


def test_kpi_delta_and_caption_render_in_table_cell():
    data = _docx_bytes(_full_payload(kpis=[
        KPISpec(label="Revenue", value="$1M", delta="+23%", caption="YoY"),
    ]))
    doc = _docx_doc(data)
    kpi_table = next(t for t in doc.tables if len(t.rows) == 2 and len(t.columns) == 1)
    cell = kpi_table.rows[1].cells[0]
    assert "$1M" in cell.text
    assert "+23%" in cell.text
    assert "YoY" in cell.text


def test_data_table_renders_chart_rows():
    data = _docx_bytes(_full_payload(chart=ChartSpec(
        title="Revenue",
        type="bar",
        x_key="quarter",
        y_keys=["revenue"],
        data=[
            {"quarter": "Q1", "revenue": 100},
            {"quarter": "Q2", "revenue": 200},
        ],
    )))
    doc = _docx_doc(data)
    # Find a table with header + 2 data rows.
    data_table = next(
        t for t in doc.tables
        if len(t.rows) == 3 and "quarter" in t.rows[0].cells[0].text.lower()
    )
    q1 = data_table.rows[1].cells[0].text
    q2 = data_table.rows[2].cells[0].text
    assert "Q1" in q1
    assert "Q2" in q2


def test_page_number_footer_present():
    """The footer should include a PAGE field (renders as "1" until Word recomputes)."""
    data = _docx_bytes(_full_payload())
    doc = _docx_doc(data)
    section = doc.sections[0]
    footer = section.footer
    footer_text = "\n".join(p.text for p in footer.paragraphs)
    assert "Generated by Zhanlu AI" in footer_text
    # The field has a placeholder "1" before Word recomputes it.
    assert "Page" in footer_text
    assert " of " in footer_text


# ---------------------------------------------------------------------------
# Edge-case tests — sparse payloads must still produce useful content
# ---------------------------------------------------------------------------


def test_no_data_payload_still_renders_cover_and_sections():
    """The exact failure mode from the user's report: no data, no
    KPIs, no insights.  Must still produce a multi-section report so
    the user opens a useful docx, not a near-empty file."""
    rcp = ReportCardPayload(
        title="sales report for me",
        source="test2",
        generated_at="2026-07-22T12:00:00",
        summary="The query executed successfully but the database contains no sales / revenue data.",
        methodology=(
            "Executed SQL: SELECT material, SUM(revenue) FROM sales GROUP BY material. "
            "Database 'test2' contains only an 'addresses' table — no sales transactions."
        ),
        kpis=[KPISpec(label="Rows returned", value="0")],
        insights=[
            InsightSpec(text="The query executed successfully but returned 0 rows."),
            InsightSpec(text="Executed SQL: SELECT material, SUM(revenue) FROM sales GROUP BY material."),
            InsightSpec(text="Try broadening the filters, date range, or search terms and run again."),
        ],
        recommendations=[
            InsightSpec(text="Connect a database with sales data, or import historical sales records."),
        ],
        next_step="Verify the database connection and try a different query.",
    )
    data = _docx_bytes(rcp)
    doc = _docx_doc(data)
    # Collect text from BOTH paragraphs and table cells (KPIs live in
    # the KPI table, not the body paragraphs).
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for t in doc.tables
        for row in t.rows
        for cell in row.cells
    )
    text = para_text + "\n" + table_text

    # Must have ALL of these sections, even with no data:
    for required in [
        "sales report for me",        # title
        "Executive Summary",          # summary section
        "Methodology",                # methodology section
        "Key Metrics",                # KPI section
        "Rows returned",              # the one KPI
        "Insights",                   # insights section
        "Recommendations",            # recommendations section
        "Connect a database",         # recommendation text
    ]:
        assert required in text, f"Missing required section: {required!r}"

    # next_step is deliberately excluded from the report — assert its absence.
    assert "Next Step" not in text
    assert "Verify the database connection" not in text


def test_legacy_minimal_payload_still_renders():
    """A payload with only the legacy fields (title + summary + kpis)
    must still render without errors and produce a docx."""
    rcp = ReportCardPayload(
        title="Legacy Report",
        summary="Just a summary.",
        kpis=[KPISpec(label="Count", value="42")],
    )
    data = _docx_bytes(rcp)
    assert data[:2] == b"PK"
    assert len(data) > 5_000


def test_payload_with_no_optional_fields_still_renders():
    """An essentially empty payload should not crash — produce a
    docx with at least the cover page."""
    rcp = ReportCardPayload(title="Empty")
    data = _docx_bytes(rcp)
    assert data[:2] == b"PK"
    doc = _docx_doc(data)
    para_text = "\n".join(p.text for p in doc.paragraphs)
    footer_text = "\n".join(
        p.text for s in doc.sections for p in s.footer.paragraphs
    )
    full_text = para_text + "\n" + footer_text
    assert "Empty" in full_text
    assert "Generated by Zhanlu AI" in full_text  # footer brand
    assert "Zhanlu AI · Intelligent Data Reports" in full_text  # cover brand


def test_render_survives_special_characters_in_text():
    """The export must HTML-escape / Word-escape `<`, `>`, `&`, `'`
    so the document opens cleanly in Word."""
    rcp = ReportCardPayload(
        title="Q3 <Review> & 'Analysis'",
        summary="Revenue & growth > 20%",
        insights=[InsightSpec(text="<script>alert(1)</script>")],
    )
    data = _docx_bytes(rcp)
    # If unescaped, the literal "<script>" in raw text would still be
    # there, but Word/python-docx should produce a valid file either
    # way.  The important assertion is that the file opens cleanly.
    doc = _docx_doc(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Q3" in text
    assert "Revenue" in text
