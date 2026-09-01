#!/usr/bin/env python3
"""Generic sandbox runner — generates file artifacts from input data.

This script runs INSIDE the Docker sandbox container. It reads the input
package from /input/ and writes the generated file to /output/.

Input layout (mounted read-only at /input/):
    /input/config.json          — {format, title, instructions, row_count}
    /input/data/*.json          — data snapshots (array of row objects)
    /input/data/*.csv           — data snapshots (CSV format)
    /input/instructions.md      — natural-language instructions
    /input/sandbox_runner.py    — this script itself (self-mount)

Output layout (writable at /output/):
    /output/report.xlsx         — generated file (extension varies by format)
    /output/build_manifest.json — metadata about what was generated

Supported formats:
    xlsx — Excel spreadsheet with auto-sized columns and header styling
    pptx — PowerPoint presentation with title slide + data table
    html — Interactive HTML dashboard with Plotly charts
    pdf  — PDF report with data table
    docx — Word document with data table
    md   — Markdown file with formatted table
"""

import json
import os
import sys
import traceback
from datetime import datetime, timezone
from pathlib import Path

INPUT_DIR = Path("/input")
DATA_DIR = INPUT_DIR / "data"
OUTPUT_DIR = Path("/output")
CONFIG_PATH = INPUT_DIR / "config.json"
INSTRUCTIONS_PATH = INPUT_DIR / "instructions.md"


def load_config():
    """Load the skill configuration from /input/config.json."""
    if not CONFIG_PATH.exists():
        print("ERROR: /input/config.json not found", file=sys.stderr)
        sys.exit(1)
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def load_data():
    """Load all data snapshots from /input/data/*.json."""
    all_rows = []
    if not DATA_DIR.exists():
        return all_rows

    for json_file in sorted(DATA_DIR.glob("*.json")):
        with open(json_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            all_rows.extend(data)
        elif isinstance(data, dict):
            all_rows.append(data)

    # Also try CSV files
    for csv_file in sorted(DATA_DIR.glob("*.csv")):
        import csv
        with open(csv_file, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                all_rows.append(dict(row))

    return all_rows


def load_instructions():
    """Load natural-language instructions from /input/instructions.md."""
    if INSTRUCTIONS_PATH.exists():
        return INSTRUCTIONS_PATH.read_text(encoding="utf-8")
    return ""


def derive_columns(rows):
    """Derive column names from the first row's keys (preserving order)."""
    if not rows:
        return []
    seen = set()
    columns = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        for key in row.keys():
            if key not in seen:
                seen.add(key)
                columns.append(str(key))
    return columns


def write_manifest(config, rows, output_file, fmt):
    """Write build_manifest.json with generation metadata."""
    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat() + "Z",
        "format": fmt,
        "title": config.get("title", "Generated Report"),
        "row_count": len(rows),
        "columns": derive_columns(rows),
        "output_file": output_file,
        "instructions": config.get("instructions", ""),
        "status": "success",
    }
    manifest_path = OUTPUT_DIR / "build_manifest.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)
    print(f"WROTE {manifest_path}")


# --- Format generators ---

def generate_xlsx(rows, config, instructions):
    """Generate an Excel spreadsheet from data rows."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = config.get("title", "Report")[:31]

    columns = derive_columns(rows)

    # Header row
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="D9D9D9"),
        right=Side(style="thin", color="D9D9D9"),
        top=Side(style="thin", color="D9D9D9"),
        bottom=Side(style="thin", color="D9D9D9"),
    )

    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # Data rows
    for row_idx, row in enumerate(rows, 2):
        for col_idx, col_name in enumerate(columns, 1):
            value = row.get(col_name, "") if isinstance(row, dict) else ""
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if row_idx % 2 == 0:
                cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

    # Auto-size columns
    for col_idx, col_name in enumerate(columns, 1):
        max_length = len(str(col_name))
        for row in rows:
            val = str(row.get(col_name, "") if isinstance(row, dict) else "")
            if len(val) > max_length:
                max_length = len(val)
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_length + 2, 50)

    # Freeze header row
    ws.freeze_panes = "A2"

    # Add summary sheet if multiple numeric columns
    if len(rows) > 0:
        ws2 = wb.create_sheet("Summary")
        ws2.cell(row=1, column=1, value="Report Title").font = Font(bold=True)
        ws2.cell(row=1, column=2, value=config.get("title", "Report"))
        ws2.cell(row=2, column=1, value="Generated At").font = Font(bold=True)
        ws2.cell(row=2, column=2, value=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
        ws2.cell(row=3, column=1, value="Total Rows").font = Font(bold=True)
        ws2.cell(row=3, column=2, value=len(rows))
        ws2.cell(row=4, column=1, value="Columns").font = Font(bold=True)
        ws2.cell(row=4, column=2, value=", ".join(columns))

    output_file = "report.xlsx"
    output_path = OUTPUT_DIR / output_file
    wb.save(str(output_path))
    print(f"WROTE {output_path}")
    return output_file


def generate_pptx(rows, config, instructions):
    """Generate a branded PowerPoint deck from data + config.

    Phase 1B: the deck is rendered by the shared ``layout_engine`` — the same
    single source of truth used by the in-process exporter (so chat-driven and
    exporter decks are byte-identical).  The orchestrator ships a plan dict in
    ``config["deck_plan"]`` (a DeckPlan.model_dump()); ``rows`` is the raw data
    snapshot used by chart / data-table slides; the brand colors live in
    ``config["theme_tokens"]`` (a hex dict produced by DeckTheme.as_hex_dict()).

    When ``config["deck_plan"]`` is absent (legacy, pre-planner sandbox path),
    a minimal plan is synthesized from the ReportCardPayload-shaped config
    fields so the old behavior still works.  ALL pptx imports are function-local
    because this script runs inside the sandbox container; ``layout_engine`` and
    ``branded_charts`` are vendored into the input package alongside this file.
    """
    deck_plan = config.get("deck_plan")
    if not deck_plan:
        # Legacy fallback: build a DeckPlan-shaped dict from the flat config.
        deck_plan = _build_plan_dict_from_config(config, rows)

    # The vendored layout_engine.render consumes a dict (it calls dict(plan)).
    # In the sandbox container the file is flat at /input/skill/layout_engine.py;
    # in repo/dev/tests it lives under app.services.artifacts.  Prefer the flat
    # vendored copy, then fall back to the dotted path so in-process tests +
    # the dispatcher parity test keep working.
    try:
        from layout_engine import render as _render_engine
    except Exception:
        from app.services.artifacts.layout_engine import render as _render_engine

    theme_tokens = config.get("theme_tokens") or {}
    if isinstance(theme_tokens, dict) and theme_tokens.get("name") is None:
        theme_tokens = dict(theme_tokens)
    style_recipe = config.get("style_recipe") or "sharp"
    ctx = {
        "theme_tokens": theme_tokens,
        "style_recipe": style_recipe,
        # Provenance label → layout_engine source_citation footer.
        "source_label": config.get("source") or "",
    }

    data = _render_engine(deck_plan, rows or [], ctx)

    output_file = "report.pptx"
    output_path = OUTPUT_DIR / output_file
    with open(output_path, "wb") as f:
        f.write(data)
    print(f"WROTE {output_path}")
    return output_file


def _build_plan_dict_from_config(config, rows):
    """Synthesize a DeckPlan-shaped dict from the legacy flat config so the
    sandbox path still works when no plan was provided by the orchestrator.
    """
    slides = []
    title = config.get("title", "Generated Report")
    source = config.get("source", "")
    summary = config.get("summary", "")
    methodology = config.get("methodology", "")
    kpis = config.get("kpis", []) or []
    insights = config.get("insights", []) or []
    key_findings = config.get("key_findings", []) or []
    recommendations = config.get("recommendations", []) or []
    sections = config.get("sections", []) or []
    next_step = config.get("next_step")
    warnings = config.get("warnings", []) or []
    chart = config.get("chart") or {}
    chart_data_rows = chart.get("data") or []

    slides.append({
        "layout": "cover",
        "title": title,
        "subtitle": (f"Source: {source}" if source else ""),
        "kpi_specs": [
            {"label": k.get("label", ""), "value": str(k.get("value", "")),
             "delta": k.get("delta"), "caption": k.get("caption")}
            for k in kpis
        ],
    })

    if sections:
        slides.append({
            "layout": "agenda",
            "title": "Agenda",
            "bullets": [
                (s.get("title", s) if isinstance(s, dict) else str(s)) for s in sections[:10]
            ],
        })

    if summary or warnings:
        slides.append({
            "layout": "insights_bullets",
            "title": "Executive Summary",
            "bullets": ([summary] if summary else [])
            + ([f"Warning: {w}" for w in warnings] if warnings else []),
        })

    if kpis:
        slides.append({
            "layout": "kpi_grid",
            "title": "Key Metrics",
            "kpi_specs": [
                {"label": k.get("label", ""), "value": str(k.get("value", "")),
                 "delta": k.get("delta"), "caption": k.get("caption")}
                for k in kpis
            ],
        })

    if chart_data_rows:
        slides.append({
            "layout": "chart_full",
            "title": chart.get("title", "Chart"),
            "chart_spec": {
                "chart_type": chart.get("type", "bar"),
                "x_key": chart.get("x_key", "label"),
                "y_keys": chart.get("y_keys") or ["value"],
                "title": chart.get("title", "Chart"),
            },
            # The layout engine renders the chart from explicit rows.
            "chart_rows": chart_data_rows,
        })

    if key_findings:
        slides.append({
            "layout": "findings_cards",
            "title": "Key Findings",
            "bullets": [
                (kf.get("text", "") if isinstance(kf, dict) else str(kf))
                for kf in key_findings[:6]
            ],
        })

    if insights:
        slides.append({
            "layout": "insights_bullets",
            "title": "Insights",
            "bullets": [
                (ins.get("text", "") if isinstance(ins, dict) else str(ins))
                for ins in insights[:6]
            ],
        })

    if recommendations:
        slides.append({
            "layout": "recommendations",
            "title": "Recommendations",
            "bullets": [
                (r.get("text", "") if isinstance(r, dict) else str(r))
                for r in recommendations[:6]
            ],
        })

    data_rows = chart_data_rows if chart_data_rows else rows
    if data_rows:
        # Use a data_table slide when rows exist; the layout engine truncates
        # to the configured top-N + highlight, matching the plan.  The builder
        # expects ``table_rows`` (row dicts) + ``table_cols`` (column keys).
        table_cols = list(data_rows[0].keys())[:10] if isinstance(data_rows[0], dict) else []
        slides.append({
            "layout": "data_table",
            "title": "Data",
            "table_rows": data_rows,
            "table_cols": table_cols,
        })

    if methodology:
        slides.append({
            "layout": "insights_bullets",
            "title": "Methodology",
            "bullets": [methodology],
        })

    if next_step:
        slides.append({
            "layout": "closing",
            "title": "Next Step",
            "bullets": [next_step],
        })

    return {
        "title": title,
        "deck_type": "data_report",
        "theme_recommendation": config.get("theme_tokens", {}).get("name", "zhanlu-blue"),
        "slides": slides,
        "summary": summary,
        "methodology": methodology,
    }


def generate_html(rows, config, instructions):
    """Generate an interactive HTML dashboard with Plotly charts.

    When the caller supplies ``config['kpis']`` (list of {label, value,
    caption} dicts) and ``config['insights']`` (list of strings or
    {text} dicts), they take precedence over the auto-derived
    "Total Records / Columns / Numeric Columns" stat cards and add an
    Insights panel below the data table.  This keeps the sidecar
    preview in lockstep with the DOCX renderer's KPI table.
    """
    columns = derive_columns(rows)
    numeric_columns = []
    for col in columns:
        numeric_count = 0
        for row in rows:
            val = row.get(col) if isinstance(row, dict) else None
            if val is not None:
                try:
                    float(val)
                    numeric_count += 1
                except (ValueError, TypeError):
                    pass
        if numeric_count > len(rows) * 0.5:
            numeric_columns.append(col)

    # Build chart specs
    charts_html = ""
    if numeric_columns and columns:
        x_col = columns[0]
        y_cols = numeric_columns[:4]  # Max 4 numeric series

        # Bar chart
        chart_data = []
        for row in rows[:50]:
            if not isinstance(row, dict):
                continue
            point = {"label": str(row.get(x_col, ""))}
            for y_col in y_cols:
                try:
                    point[y_col] = float(row.get(y_col, 0))
                except (ValueError, TypeError):
                    point[y_col] = 0
            chart_data.append(point)

        import json as _json
        chart_data_json = _json.dumps(chart_data)

        charts_html = f"""
        <div id="chart_div"></div>
        <script>
            var chartData = {chart_data_json};
            var trace1 = {{
                x: chartData.map(function(d) {{ return d.label; }}),
                y: chartData.map(function(d) {{ return d["{y_cols[0]}"] || 0; }}),
                type: 'bar',
                name: '{y_cols[0]}',
                marker: {{ color: '#4472C4' }}
            }};
            var data = [trace1];
            var layout = {{
                title: '{y_cols[0]} by {x_col}',
                xaxis: {{ tickangle: -45 }},
                margin: {{ t: 40, b: 100 }}
            }};
            Plotly.newPlot('chart_div', data, layout, {{responsive: true}});
        </script>
        """

    # Build data table HTML
    table_rows_html = ""
    for row in rows[:100]:
        if not isinstance(row, dict):
            continue
        cells = ""
        for col in columns:
            val = row.get(col, "")
            cells += f"<td>{val}</td>"
        table_rows_html += f"<tr>{cells}</tr>"

    header_cells = "".join(f"<th>{col}</th>" for col in columns)

    # --- KPI cards (prefer caller-supplied; fall back to defaults) ---
    kpis = config.get("kpis") or []
    if kpis:
        kpi_cards = "".join(
            f'<div class="stat-card">'
            f'<div class="label">{k.get("label", "")}</div>'
            f'<div class="value">{k.get("value", "")}</div>'
            + (f'<div class="caption">{k.get("caption", "")}</div>' if k.get("caption") else "")
            + '</div>'
            for k in kpis
        )
    else:
        kpi_cards = (
            f'<div class="stat-card"><div class="label">Total Records</div><div class="value">{len(rows)}</div></div>'
            f'<div class="stat-card"><div class="label">Columns</div><div class="value">{len(columns)}</div></div>'
            f'<div class="stat-card"><div class="label">Numeric Columns</div><div class="value">{len(numeric_columns)}</div></div>'
        )

    # --- Insights panel ---
    insights = config.get("insights") or []
    insights_section = ""
    if insights:
        items = "".join(
            f"<li>{ins if isinstance(ins, str) else (ins.get('text', '') if isinstance(ins, dict) else '')}</li>"
            for ins in insights
        )
        insights_section = f"""
        <div class="insights-container">
            <h2>Insights</h2>
            <ul>{items}</ul>
        </div>
        """

    # --- Key findings (narrative paragraphs) ---
    key_findings = config.get("key_findings") or []
    key_findings_section = ""
    if key_findings:
        paras = "".join(
            f"<p>{f if isinstance(f, str) else (f.get('text', '') if isinstance(f, dict) else '')}</p>"
            for f in key_findings
        )
        key_findings_section = f"""
        <div class="findings-container">
            <h2>Key Findings</h2>
            {paras}
        </div>
        """

    # --- Recommendations (bullets) ---
    recommendations = config.get("recommendations") or []
    recommendations_section = ""
    if recommendations:
        items = "".join(
            f"<li>{r if isinstance(r, str) else (r.get('text', '') if isinstance(r, dict) else '')}</li>"
            for r in recommendations
        )
        recommendations_section = f"""
        <div class="insights-container">
            <h2>Recommendations</h2>
            <ul>{items}</ul>
        </div>
        """

    # --- Custom sections (heading + paragraph/bullets) ---
    custom_sections = ""
    for sec in (config.get("sections") or []):
        if not isinstance(sec, dict) or not sec.get("title"):
            continue
        body = ""
        if sec.get("content"):
            body += f"<p>{sec['content']}</p>"
        for b in (sec.get("bullets") or []):
            body += f"<ul><li>{b}</li></ul>"
        custom_sections += f"""
        <div class="insights-container">
            <h2>{sec['title']}</h2>
            {body}
        </div>
        """

    # --- Executive summary callout ---
    summary = (config.get("summary") or "").strip()
    summary_section = ""
    if summary:
        summary_section = f"""
        <div class="summary-container">
            <span class="badge">EXECUTIVE SUMMARY</span>
            <p>{summary}</p>
        </div>
        """

    # --- Methodology callout ---
    methodology = (config.get("methodology") or "").strip()
    methodology_section = ""
    if methodology:
        methodology_section = f"""
        <div class="methodology-container">
            <span class="badge subtle">METHODOLOGY</span>
            <p>{methodology}</p>
        </div>
        """

    # --- Next step callout ---
    next_step = (config.get("next_step") or "").strip()
    next_step_section = ""
    if next_step:
        next_step_section = f"""
        <div class="next-step-container">
            <strong>Next step:</strong> {next_step}
        </div>
        """

    # --- SQL block ---
    sql_text = (config.get("sql") or "").strip()
    sql_section = ""
    if sql_text:
        sql_section = f"""
        <div class="insights-container">
            <h2>SQL</h2>
            <pre style="background:#0F172A;color:#E2E8F0;padding:12px;border-radius:6px;
                        overflow-x:auto;font-size:12px;">{sql_text}</pre>
        </div>
        """

    # Full HTML document
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{config.get('title', 'Dashboard')}</title>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f5f5f5; color: #333; }}
        .header {{ background: linear-gradient(135deg, #1f2a44 0%, #e67e22 100%); color: white; padding: 30px 40px; }}
        .header h1 {{ font-size: 28px; margin-bottom: 8px; }}
        .header .meta {{ font-size: 14px; opacity: 0.9; }}
        .container {{ max-width: 1200px; margin: 0 auto; padding: 24px; }}
        .stats {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 16px; margin-bottom: 24px; }}
        .stat-card {{ background: white; border-radius: 12px; padding: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .stat-card .label {{ font-size: 12px; text-transform: uppercase; letter-spacing: 0.5px; color: #888; margin-bottom: 4px; }}
        .stat-card .value {{ font-size: 28px; font-weight: 700; color: #1f2a44; }}
        .stat-card .caption {{ font-size: 11px; color: #94A3B8; margin-top: 4px; }}
        .chart-container {{ background: white; border-radius: 12px; padding: 24px; margin-bottom: 24px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .table-container {{ background: white; border-radius: 12px; padding: 24px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); overflow-x: auto; }}
        .insights-container {{ background: white; border-radius: 12px; padding: 24px; margin-top: 24px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .insights-container h2 {{ color: #e67e22; font-size: 18px; margin-bottom: 12px; }}
        .insights-container ul {{ list-style: disc; padding-left: 24px; }}
        .insights-container li {{ margin-bottom: 8px; line-height: 1.5; }}
        .findings-container {{ background: white; border-radius: 12px; padding: 24px; margin-top: 24px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .findings-container h2 {{ color: #1f2a44; font-size: 18px; margin-bottom: 12px; }}
        .findings-container p {{ line-height: 1.65; margin: 10px 0; color: #1E293B; }}
        .summary-container {{ background: #FEF3C7; border-left: 4px solid #F59E0B; border-radius: 6px; padding: 16px 20px; margin-bottom: 24px; }}
        .summary-container p {{ margin: 6px 0 0; line-height: 1.55; }}
        .methodology-container {{ background: #F1F5F9; border-left: 4px solid #64748B; border-radius: 6px; padding: 16px 20px; margin-bottom: 24px; }}
        .methodology-container p {{ margin: 6px 0 0; line-height: 1.55; }}
        .next-step-container {{ background: #EFF6FF; border-left: 4px solid #2563EB; border-radius: 6px; padding: 12px 16px; margin-bottom: 16px; }}
        .badge {{ display: inline-block; background: #92400E; color: #fff; font-size: 10px; font-weight: 700; letter-spacing: 1px; padding: 3px 8px; border-radius: 4px; margin-bottom: 6px; }}
        .badge.subtle {{ background: #475569; }}
        table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
        th {{ background: #1f2a44; color: white; padding: 10px 12px; text-align: left; font-weight: 600; position: sticky; top: 0; }}
        td {{ padding: 8px 12px; border-bottom: 1px solid #eee; }}
        tr:nth-child(even) {{ background: #fafafa; }}
        tr:hover {{ background: #f0f7ff; }}
        .footer {{ text-align: center; padding: 20px; color: #888; font-size: 12px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{config.get('title', 'Dashboard')}</h1>
        <div class="meta">Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | {len(rows)} records | {len(columns)} columns</div>
    </div>
    <div class="container">
        {summary_section}
        {methodology_section}
        <div class="stats">{kpi_cards}</div>
        <div class="chart-container">
            {charts_html if charts_html else '<p style="text-align:center;color:#888;padding:40px;">No numeric data available for charts</p>'}
        </div>
        <div class="table-container">
            <table>
                <thead><tr>{header_cells}</tr></thead>
                <tbody>{table_rows_html}</tbody>
            </table>
        </div>
        {key_findings_section}
        {insights_section}
        {recommendations_section}
        {custom_sections}
        {sql_section}
        {next_step_section}
    </div>
    <div class="footer">Generated by Zhanlu AI</div>
</body>
</html>"""

    output_file = "report.html"
    output_path = OUTPUT_DIR / output_file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"WROTE {output_path}")
    return output_file


def generate_pdf(rows, config, instructions):
    """Generate a PDF report from data rows."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    output_file = "report.pdf"
    output_path = OUTPUT_DIR / output_file

    doc = SimpleDocTemplate(str(output_path), pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    elements = []

    # Title
    title_style = styles["Heading1"]
    elements.append(Paragraph(config.get("title", "Report"), title_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Meta info
    meta_style = styles["Normal"]
    elements.append(Paragraph(
        f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | {len(rows)} records",
        meta_style,
    ))
    elements.append(Spacer(1, 0.3 * inch))

    if instructions:
        elements.append(Paragraph(f"<b>Instructions:</b> {instructions}", meta_style))
        elements.append(Spacer(1, 0.2 * inch))

    # Data table
    columns = derive_columns(rows)
    table_data = [columns]
    for row in rows[:100]:
        if not isinstance(row, dict):
            continue
        table_data.append([str(row.get(col, "")) for col in columns])

    table = Table(table_data)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9D9D9")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
    ]))

    elements.append(table)
    doc.build(elements)
    print(f"WROTE {output_path}")
    return output_file


def generate_docx(rows, config, instructions):
    """Generate a Claude-style Word document from data rows.

    Produces a richly structured report (matches the visual language of
    the rich-HTML dashboard generated by ``generate_html``):

    1. Cover page — themed title, source, generation date, brand line
    2. Executive Summary — config['summary'] in a tinted callout
    3. Methodology — config['methodology'] (if present)
    4. Key Metrics — config['kpis'] as a 1×N table with bold values
    5. Key Findings — config['key_findings'] as narrative paragraphs
    6. Insights — config['insights'] as a bulleted list
    7. Recommendations — config['recommendations'] as a bulleted list
    8. Custom sections — config['sections'] in declared order
    9. Data table — colored header row, banded rows
    10. SQL block — config['sql'] in a tinted code paragraph
    11. Next Step — config['next_step'] as a tinted callout
    12. Footer with page numbers — "Generated by Zhanlu AI · Page X of Y"

    Proper Word styles ("Heading 1", "Heading 2") are used so mammoth's
    docx→html conversion preserves structure.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Theme — kept consistent with generate_html's CSS palette so the
    # user gets the same brand feel in both formats.
    ORANGE = RGBColor(0xE6, 0x7E, 0x22)   # accent / section headings
    DARK = RGBColor(0x1F, 0x2A, 0x44)     # primary text / KPI values
    MUTED = RGBColor(0x6B, 0x72, 0x80)    # captions / labels
    HEADER_FILL = "1F2A44"                # dark navy header background

    doc = Document()

    # Set default body font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # Style heading 1/2/3 in dark navy so mammoth converts them to
    # <h1>/<h2>/<h3> with a consistent look.
    for level, (size, color) in {
        1: (16, DARK),
        2: (14, ORANGE),
        3: (12, MUTED),
    }.items():
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = color
        h_style.font.bold = True
        h_style.font.name = "Calibri"

    # --- 1. Cover page ---
    title_text = config.get("title", "Report")
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover_title.add_run(f"📊  {title_text}")
    cover_run.font.size = Pt(28)
    cover_run.font.bold = True
    cover_run.font.color.rgb = DARK
    cover_title.paragraph_format.space_before = Pt(120)
    cover_title.paragraph_format.space_after = Pt(12)

    source = (config.get("source") or "").strip()
    if source:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(f"Source: {source}")
        sr.font.size = Pt(13)
        sr.font.color.rgb = MUTED

    gen_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(2)
    sr2 = sub2.add_run(f"Generated {gen_at} · {len(rows)} record{'s' if len(rows) != 1 else ''}")
    sr2.font.size = Pt(11)
    sr2.font.color.rgb = MUTED

    # Brand footer of cover page
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(180)
    br = brand.add_run("Zhanlu AI · Intelligent Data Reports")
    br.font.size = Pt(10)
    br.font.color.rgb = MUTED
    br.italic = True

    # Page break to start the body
    doc.add_page_break()

    def _section_heading(text, level=1):
        """Add a styled section heading."""
        p = doc.add_heading(text, level=level)
        if level == 1:
            for r in p.runs:
                r.font.color.rgb = DARK
                r.font.size = Pt(16)
        elif level == 2:
            for r in p.runs:
                r.font.color.rgb = ORANGE
                r.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _shade_cell(cell, hex_color):
        """Set a cell's background color (hex string without #)."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _shade_paragraph(paragraph, hex_color):
        """Set a paragraph's background color (for tinted callouts)."""
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    def _insight_text(it):
        if isinstance(it, str):
            return it
        if isinstance(it, dict):
            return it.get("text", "")
        return str(it) if it is not None else ""

    def _add_insight_bullets(items):
        for ins in items or []:
            text = _insight_text(ins)
            if not text:
                continue
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(10)

    # --- 2. Executive Summary ---
    summary = (config.get("summary") or "").strip()
    if summary:
        _section_heading("Executive Summary")
        para = doc.add_paragraph(summary)
        _shade_paragraph(para, "FEF3C7")
        para.paragraph_format.left_indent = Cm(0.4)
        para.paragraph_format.space_after = Pt(12)
        for r in para.runs:
            r.font.size = Pt(10)

    # --- 3. Methodology ---
    methodology = (config.get("methodology") or "").strip()
    if methodology:
        _section_heading("Methodology")
        para = doc.add_paragraph(methodology)
        _shade_paragraph(para, "F1F5F9")
        para.paragraph_format.left_indent = Cm(0.4)
        for r in para.runs:
            r.font.size = Pt(10)

    # --- 4. KPI table ---
    kpis = config.get("kpis") or []
    if kpis:
        _section_heading("Key Metrics")
        kpi_table = doc.add_table(rows=2, cols=len(kpis))
        kpi_table.style = "Light Grid Accent 1"
        # Header row — labels
        for i, kpi in enumerate(kpis):
            cell = kpi_table.rows[0].cells[i]
            cell.text = ""
            _shade_cell(cell, "F8FAFC")
            label_p = cell.paragraphs[0]
            label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = label_p.add_run(str(kpi.get("label", "")))
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = MUTED
            label_run.font.bold = True
        # Value row
        for i, kpi in enumerate(kpis):
            cell = kpi_table.rows[1].cells[i]
            cell.text = ""
            value_p = cell.paragraphs[0]
            value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value_run = value_p.add_run(str(kpi.get("value", "")))
            value_run.font.size = Pt(22)
            value_run.font.bold = True
            value_run.font.color.rgb = DARK
            delta = kpi.get("delta")
            if delta:
                dp = cell.add_paragraph()
                dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dr = dp.add_run(str(delta))
                dr.font.size = Pt(9)
                dr.font.bold = True
                dr.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            caption = kpi.get("caption")
            if caption:
                cap_p = cell.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(str(caption))
                cap_run.font.size = Pt(8)
                cap_run.font.color.rgb = MUTED
                cap_run.italic = True
        doc.add_paragraph()  # spacing after the KPI row

    # --- 5. Key Findings (narrative paragraphs) ---
    key_findings = config.get("key_findings") or []
    if key_findings:
        _section_heading("Key Findings")
        for f in key_findings:
            text = _insight_text(f)
            if not text:
                continue
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.size = Pt(10)

    # --- 6. Insights (bullets) ---
    insights = config.get("insights") or []
    if insights:
        _section_heading("Insights")
        _add_insight_bullets(insights)

    # --- 7. Recommendations (bullets) ---
    recommendations = config.get("recommendations") or []
    if recommendations:
        _section_heading("Recommendations")
        _add_insight_bullets(recommendations)

    # --- 8. Custom sections ---
    sections = config.get("sections") or []
    for sec in sections:
        if not isinstance(sec, dict) or not sec.get("title"):
            continue
        _section_heading(sec.get("title"))
        if sec.get("content"):
            p = doc.add_paragraph(sec["content"])
            for r in p.runs:
                r.font.size = Pt(10)
        for b in (sec.get("bullets") or []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(b))
            run.font.size = Pt(10)

    # --- Instructions (legacy field, kept for backwards compat) ---
    if instructions and not (summary or methodology or key_findings or insights):
        _section_heading("Instructions")
        p = doc.add_paragraph(instructions)
        for r in p.runs:
            r.font.size = Pt(10)

    # --- 9. Data table ---
    columns = derive_columns(rows)
    if columns and rows:
        _section_heading("Data")
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Light Grid Accent 1"

        # Header row — dark fill, white bold text
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(columns):
            cell = header_cells[i]
            _shade_cell(cell, HEADER_FILL)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(col_name))
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

        # Data rows — banded for readability
        for row_idx, row in enumerate(rows[:100]):
            if not isinstance(row, dict):
                continue
            row_cells = table.add_row().cells
            banded = (row_idx % 2 == 1)
            for i, col_name in enumerate(columns):
                cell = row_cells[i]
                cell.text = str(row.get(col_name, ""))
                if banded:
                    _shade_cell(cell, "F2F2F2")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

    # --- 10. SQL block ---
    sql = (config.get("sql") or "").strip()
    if sql:
        _section_heading("SQL")
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(sql)
        sql_run.font.name = "Consolas"
        sql_run.font.size = Pt(9)
        sql_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        _shade_paragraph(sql_para, "F1F5F9")

    # --- 11. Next step ---
    next_step = (config.get("next_step") or "").strip()
    if next_step:
        _section_heading("Next Step")
        ns = doc.add_paragraph(next_step)
        _shade_paragraph(ns, "EFF6FF")
        ns.paragraph_format.left_indent = Cm(0.4)
        for r in ns.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x25, 0x6, 0xEB)
            r.font.size = Pt(10)

    # --- 12. Footer with page numbers ---
    _add_page_number_footer(doc, "Generated by Zhanlu AI")

    output_file = "report.docx"
    output_path = OUTPUT_DIR / output_file
    doc.save(str(output_path))
    print(f"WROTE {output_path}")
    return output_file


def _add_page_number_footer(doc, brand_text: str) -> None:
    """Add a "brand · Page X of Y" footer to every page using
    Word's PAGE and NUMPAGES fields."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    section = doc.sections[0]
    footer = section.footer
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    muted = RGBColor(0x94, 0xA3, 0xB8)
    # "brand_text · Page "
    run = para.add_run(f"{brand_text} · Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = muted

    _add_field(para, "PAGE")
    para.add_run(" of ").font.size = Pt(9)
    _add_field(para, "NUMPAGES")
    for r in para.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = muted


def _add_field(paragraph, field_code: str) -> None:
    """Append a Word field (e.g. 'PAGE', 'NUMPAGES') to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_code} "
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)


def generate_md(rows, config, instructions):
    """Generate a Markdown file from data rows."""
    columns = derive_columns(rows)

    lines = [
        f"# {config.get('title', 'Report')}",
        "",
        f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | {len(rows)} records",
        "",
    ]

    if instructions:
        lines.extend(["## Instructions", "", instructions, ""])

    lines.extend(["## Data", ""])

    # Markdown table
    header = "| " + " | ".join(columns) + " |"
    separator = "| " + " | ".join(["---"] * len(columns)) + " |"
    lines.append(header)
    lines.append(separator)

    for row in rows[:100]:
        if not isinstance(row, dict):
            continue
        cells = [str(row.get(col, "")) for col in columns]
        lines.append("| " + " | ".join(cells) + " |")

    lines.append("")
    lines.append("---")
    lines.append("*Generated by Zhanlu Sandbox Runner*")

    output_file = "report.md"
    output_path = OUTPUT_DIR / output_file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"WROTE {output_path}")
    return output_file


# --- Main ---

GENERATORS = {
    "xlsx": generate_xlsx,
    "pptx": generate_pptx,
    "html": generate_html,
    "pdf":  generate_pdf,
    "docx": generate_docx,
    "md":   generate_md,
}


def main():
    """Entry point — read input, generate output, write manifest."""
    print("=== Zhanlu Sandbox Runner ===")
    print(f"Input dir: {INPUT_DIR}")
    print(f"Output dir: {OUTPUT_DIR}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load inputs
    config = load_config()
    rows = load_data()
    instructions = load_instructions()

    fmt = config.get("format", "xlsx").lower()
    print(f"Format: {fmt}")
    print(f"Title: {config.get('title', 'N/A')}")
    print(f"Rows loaded: {len(rows)}")

    if not rows:
        print("WARNING: No data rows loaded — generating empty report")

    # Generate
    generator = GENERATORS.get(fmt)
    if not generator:
        print(f"ERROR: Unknown format '{fmt}'", file=sys.stderr)
        sys.exit(1)

    try:
        output_file = generator(rows, config, instructions)
        write_manifest(config, rows, output_file, fmt)
        print(f"SUCCESS: Generated {output_file}")
    except Exception as e:
        print(f"ERROR: Generation failed: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        # Write error manifest
        error_manifest = {
            "generated_at": datetime.now(timezone.utc).isoformat() + "Z",
            "format": fmt,
            "title": config.get("title", "Report"),
            "status": "failed",
            "error": str(e),
        }
        with open(OUTPUT_DIR / "build_manifest.json", "w") as f:
            json.dump(error_manifest, f, indent=2)
        sys.exit(1)


if __name__ == "__main__":
    main()
