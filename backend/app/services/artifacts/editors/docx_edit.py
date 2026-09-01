"""DOCX round-trip edit operations.

Supported ops:
* ``replace_text`` — ``{"op": "replace_text", "find": ..., "replace": ...}``
"""

from __future__ import annotations

import io

from docx import Document

from app.services.artifacts.editors import EditError


def apply_docx_edits(data: bytes, operations: list[dict]) -> tuple[bytes, list[str]]:
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise EditError(f"Could not open docx: {e}") from e

    changelog: list[str] = []

    for op in operations:
        if not isinstance(op, dict):
            raise EditError(f"Invalid operation: {op!r}")
        kind = op.get("op")
        if kind == "replace_text":
            _replace_text(doc, op)
            changelog.append(
                f"replace_text: {op.get('find')!r} -> {op.get('replace')!r}"
            )
        else:
            raise EditError(f"Unknown docx edit op: {kind}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), changelog


def _iter_paragraphs(doc):
    """Yield every paragraph (body + table cells)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_text(doc, op: dict) -> int:
    find = op.get("find")
    replace = op.get("replace", "")
    if not find:
        raise EditError("replace_text requires a non-empty 'find'")

    found = 0
    for p in _iter_paragraphs(doc):
        for run in p.runs:
            if find in run.text:
                run.text = run.text.replace(find, replace)
                found += 1

    if found == 0:
        raise EditError(f"Text '{find}' not found")
    return found
