"""Deterministic repairs for DOCX documents (P0.1 self-healing loop).

* ``body_font`` — raise any run below the 10pt body floor to 10pt
  (covers both body paragraphs and table cells).
"""

from __future__ import annotations

import io
import logging
from typing import Optional

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

BODY_FONT_FLOOR_PT = 10.0

SUPPORTED_RULES = {"body_font"}


def repair_doc(data: bytes, rule_ids: set[str]) -> Optional[bytes]:
    """Apply deterministic repairs for the given FAIL rule ids.

    Returns the repaired DOCX bytes, or ``None`` when nothing is fixable.
    """
    if not data:
        return None

    wanted = set(rule_ids or []) & SUPPORTED_RULES
    if not wanted:
        return None

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001 — corrupted bytes must not crash the loop
        logger.warning("repair_doc: failed to open docx: %s", e)
        return None

    changed = False
    if "body_font" in wanted:
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.size is not None and r.font.size.pt < BODY_FONT_FLOOR_PT:
                    r.font.size = Pt(BODY_FONT_FLOOR_PT)
                    changed = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.font.size is not None and r.font.size.pt < BODY_FONT_FLOOR_PT:
                                r.font.size = Pt(BODY_FONT_FLOOR_PT)
                                changed = True

    if not changed:
        return data

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
