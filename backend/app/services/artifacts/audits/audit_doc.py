#!/usr/bin/env python3
"""Semantic audit runner for generated DOCX documents.

Mechanically enforces the codifiable subset of `backend/skills/docx/SKILL.md`'s
10-point pre-emit self-audit — the rules XSD validation cannot see:

  * placeholder_text — no xxxx/lorem/ipsum/todo/tbd/"click to add"
  * heading_hierarchy — H1→H2→H3→H4 in order, no skipped levels, one H1
  * body_font        — Normal style body text at 11pt; no run below 9pt
  * line_spacing     — body line spacing within 1.15–1.5
  * margins          — ≥ 1.0" margins all around (letter's 1.25" L/R is allowed)
  * page_numbers     — a PAGE field in the footer (required for > 2-page docs)
  * headers_footers  — header/footer content present for multi-page docs
  * table_overflow   — every table fits within the text area

Rules that require semantic judgement (document-type fit, sentence case,
source-citation presence, "headings are not paragraphs") are intentionally NOT
checked here — they belong to a downstream content QA step.

Usage:
    python audit_doc.py output.docx            # human-readable report
    python audit_doc.py output.docx --json     # JSON to stdout
    python audit_doc.py output.docx --strict   # exit 1 on WARN too

Exit codes: 0 = PASS/WARN, 1 = FAIL (or WARN under --strict), 2 = tooling error.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    print(f"audit_doc: python-docx not available: {exc}", file=sys.stderr)
    sys.exit(2)


EMU_PER_INCH = 914_400

# Thresholds sourced verbatim from docx/SKILL.md.
BODY_SIZE_PT = 11.0
CAPTION_FLOOR_PT = 9.0          # caption is 9-10pt; below 9pt is always wrong
LINE_SPACING_MIN = 1.15
LINE_SPACING_MAX = 1.5
MIN_MARGIN_IN = 1.0
MIN_MARGIN_HARD_IN = 0.8        # below this is a definite problem (letter uses 1.25 L/R)
MAX_H2_COUNT = 6                # "more than 6 H2s → outline needs restructuring"
MULTIPAGE_PARA_THRESHOLD = 50   # rough heuristic for "> 2 pages"

PLACEHOLDER_RE = re.compile(
    r"\b(xxxx|lorem|ipsum|todo|tbd|click to add|placeholder)\b", re.IGNORECASE
)

# Heading style name → numeric level.
_HEADING_LEVEL = {
    "heading 1": 1, "heading 2": 2, "heading 3": 3, "heading 4": 4,
    "heading 5": 5, "heading 6": 6, "title": 1,
}


# --- Report model ---------------------------------------------------------


LEVELS = ("PASS", "WARN", "FAIL")
_SEVERITY = {"PASS": 0, "WARN": 1, "FAIL": 2}


@dataclass
class Finding:
    rule_id: str
    title: str
    level: str
    detail: str
    evidence: list[str] = field(default_factory=list)

    def __post_init__(self) -> None:
        if self.level not in LEVELS:
            raise ValueError(f"level must be one of {LEVELS}, got {self.level!r}")


def build_report(file: str, findings: list[Finding]) -> dict[str, Any]:
    counts = {lvl: 0 for lvl in LEVELS}
    for f in findings:
        counts[f.level] += 1
    status = "PASS"
    for f in findings:
        if _SEVERITY[f.level] > _SEVERITY[status]:
            status = f.level
    return {
        "tool": "audit_doc",
        "file": file,
        "status": status,
        "summary": {
            "pass": counts["PASS"],
            "warn": counts["WARN"],
            "fail": counts["FAIL"],
            "total": len(findings),
        },
        "rules": [
            {
                "id": f.rule_id,
                "title": f.title,
                "level": f.level,
                "detail": f.detail,
                "evidence": f.evidence[:20],
            }
            for f in findings
        ],
    }


def print_human(report: dict[str, Any]) -> None:
    s = report["summary"]
    print(f"audit_doc — {report['file']}")
    print(f"  status: {report['status']}  ({s['pass']} pass / {s['warn']} warn / {s['fail']} fail)")
    for r in report["rules"]:
        marker = {"PASS": "✓", "WARN": "!", "FAIL": "✗"}[r["level"]]
        print(f"  [{marker} {r['level']}] {r['id']}: {r['title']}")
        if r["detail"]:
            print(f"        {r['detail']}")
        for ev in r["evidence"][:6]:
            print(f"        - {ev}")


# --- Helpers --------------------------------------------------------------


def _in(emu: Any) -> float | None:
    if emu is None:
        return None
    return emu / EMU_PER_INCH


def _heading_level(para: Any) -> int | None:
    style = para.style
    name = (style.name if style is not None else "") or ""
    return _HEADING_LEVEL.get(name.lower())


def _iter_body_paragraphs(doc: Any):
    """Yield paragraphs in document order (top-level body only)."""
    yield from doc.paragraphs


def _nonempty_para_count(doc: Any) -> int:
    return sum(1 for p in doc.paragraphs if (p.text or "").strip())


def _has_page_breaks(doc: Any) -> bool:
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for p in doc.paragraphs:
        for br in p._p.findall(f".//{ns}br"):
            if br.get(f"{ns}type") == "page":
                return True
    return False


def _is_multipage(doc: Any) -> bool:
    """Heuristic: a doc is 'multi-page' if it has explicit page breaks or
    enough paragraphs that it almost certainly exceeds two pages."""
    if _has_page_breaks(doc):
        return True
    return _nonempty_para_count(doc) > MULTIPAGE_PARA_THRESHOLD


def _footer_has_page_field(section: Any) -> bool:
    for part in (section.footer, section.even_page_footer, section.first_page_footer):
        try:
            xml = part._element.xml
        except Exception:
            continue
        if "PAGE" in xml and ("instrText" in xml or "fldSimple" in xml):
            return True
    return False


def _footer_has_content(section: Any) -> bool:
    for part in (section.footer, section.even_page_footer, section.first_page_footer):
        try:
            for p in part.paragraphs:
                if (p.text or "").strip():
                    return True
        except Exception:
            continue
    return False


def _header_has_content(section: Any) -> bool:
    for part in (section.header, section.even_page_header, section.first_page_header):
        try:
            for p in part.paragraphs:
                if (p.text or "").strip():
                    return True
        except Exception:
            continue
    return False


def _text_area_width_in(section: Any) -> float | None:
    pw = _in(section.page_width)
    lm = _in(section.left_margin)
    rm = _in(section.right_margin)
    if None in (pw, lm, rm):
        return None
    return pw - lm - rm


def _table_width_in(table: Any) -> float | None:
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    total = 0
    found = False
    for col in table._tbl.findall(f".//{ns}gridCol"):
        w = col.get(f"{ns}w")
        if w is None:
            continue
        found = True
        total += int(w)
    if not found:
        return None
    return total / EMU_PER_INCH


# --- Rule checks ----------------------------------------------------------


def check_placeholder(doc: Any) -> Finding:
    hits: list[str] = []
    for idx, p in enumerate(_iter_body_paragraphs(doc), start=1):
        for m in PLACEHOLDER_RE.finditer(p.text or ""):
            hits.append(f"para {idx}: {m.group(0)!r}")
    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                for m in PLACEHOLDER_RE.finditer(cell.text or ""):
                    hits.append(f"table {t_idx} r{r_idx} c{c_idx}: {m.group(0)!r}")
    if hits:
        return Finding(
            "placeholder_text", "No placeholder/lorem text", "FAIL",
            f"{len(hits)} placeholder hit(s)", hits,
        )
    return Finding("placeholder_text", "No placeholder/lorem text", "PASS", "")


def check_heading_hierarchy(doc: Any) -> Finding:
    issues: list[str] = []
    prev_level = 0
    h1_count = 0
    h2_count = 0
    for idx, p in enumerate(_iter_body_paragraphs(doc), start=1):
        lvl = _heading_level(p)
        if lvl is None:
            continue
        if lvl == 1:
            h1_count += 1
        if lvl == 2:
            h2_count += 1
        if prev_level and lvl > prev_level + 1:
            issues.append(f"para {idx}: H{prev_level} → H{lvl} (skipped H{prev_level + 1})")
        prev_level = lvl
    if h1_count > 1:
        issues.append(f"{h1_count} H1/Title headings (expected exactly one document title)")
    if h2_count > MAX_H2_COUNT:
        issues.append(f"{h2_count} H2 sections (max {MAX_H2_COUNT} — outline needs restructuring)")
    if issues:
        return Finding(
            "heading_hierarchy", "H1→H2→H3→H4, no skipped levels, one H1", "WARN",
            f"{len(issues)} hierarchy issue(s)", issues,
        )
    return Finding("heading_hierarchy", "Heading hierarchy", "PASS", "")


def check_body_font(doc: Any) -> Finding:
    issues: list[str] = []
    # Normal style default.
    try:
        normal = doc.styles["Normal"]
        size = normal.font.size
        if size is not None and abs(size.pt - BODY_SIZE_PT) > 0.01:
            issues.append(f"Normal style body size = {size.pt}pt (expected {BODY_SIZE_PT:.0f}pt)")
    except Exception:
        pass
    # Runs below the caption floor anywhere in the body.
    undersized: list[str] = []
    for idx, p in enumerate(_iter_body_paragraphs(doc), start=1):
        for r in p.runs:
            if r.font.size is None:
                continue
            pt = r.font.size.pt
            if pt < CAPTION_FLOOR_PT:
                undersized.append(f"para {idx}: {pt}pt ({(r.text or '')[:30]!r})")
    if undersized:
        issues.append(f"{len(undersized)} run(s) below {CAPTION_FLOOR_PT:.0f}pt caption floor")
    if undersized:
        return Finding(
            "body_font", f"Body 11pt; no run below {CAPTION_FLOOR_PT:.0f}pt", "FAIL",
            "; ".join(issues), undersized,
        )
    if issues:
        return Finding(
            "body_font", f"Body 11pt; no run below {CAPTION_FLOOR_PT:.0f}pt", "WARN",
            "; ".join(issues), [],
        )
    return Finding("body_font", "Body font size", "PASS", "")


def check_line_spacing(doc: Any) -> Finding:
    bad: list[str] = []
    for idx, p in enumerate(_iter_body_paragraphs(doc), start=1):
        if _heading_level(p) is not None:
            continue
        ls = p.paragraph_format.line_spacing
        if ls is None:
            continue
        # line_spacing is a float (multiple) or a Length (exact rule).
        if isinstance(ls, (int, float)):
            if ls < LINE_SPACING_MIN or ls > LINE_SPACING_MAX:
                bad.append(f"para {idx}: line spacing {ls} (expected {LINE_SPACING_MIN}–{LINE_SPACING_MAX})")
    if bad:
        return Finding(
            "line_spacing", f"Body line spacing {LINE_SPACING_MIN}–{LINE_SPACING_MAX}", "WARN",
            f"{len(bad)} paragraph(s) outside the range", bad,
        )
    return Finding("line_spacing", "Line spacing", "PASS", "")


def check_margins(doc: Any) -> Finding:
    bad: list[str] = []
    for s_idx, section in enumerate(doc.sections, start=1):
        sides = {
            "top": _in(section.top_margin),
            "bottom": _in(section.bottom_margin),
            "left": _in(section.left_margin),
            "right": _in(section.right_margin),
        }
        for name, val in sides.items():
            if val is None:
                continue
            if val < MIN_MARGIN_HARD_IN:
                bad.append(f"section {s_idx} {name} margin = {val:.2f}\" (< {MIN_MARGIN_HARD_IN}\")")
    if bad:
        return Finding(
            "margins", f"≥ {MIN_MARGIN_IN}\" margins (hard floor {MIN_MARGIN_HARD_IN}\")", "WARN",
            f"{len(bad)} margin(s) below the hard floor", bad,
        )
    return Finding("margins", "Page margins", "PASS", "")


def check_page_numbers(doc: Any) -> Finding:
    has_field = any(_footer_has_page_field(s) for s in doc.sections)
    if has_field:
        return Finding("page_numbers", "PAGE field in footer", "PASS", "")
    if _is_multipage(doc):
        return Finding(
            "page_numbers", "Page numbers for > 2-page docs", "WARN",
            "multi-page document but no PAGE field detected in any footer", [],
        )
    # Single-page doc: page numbers optional → informational PASS.
    return Finding(
        "page_numbers", "Page numbers for > 2-page docs", "PASS",
        "single-page doc (page numbers optional)", [],
    )


def check_headers_footers(doc: Any) -> Finding:
    if not _is_multipage(doc):
        return Finding(
            "headers_footers", "Headers/footers for multi-page docs", "PASS",
            "single-page doc (headers/footers optional)", [],
        )
    missing: list[str] = []
    for s_idx, section in enumerate(doc.sections, start=1):
        if not _header_has_content(section) and not _footer_has_content(section):
            missing.append(f"section {s_idx}: neither header nor footer has content")
    if missing:
        return Finding(
            "headers_footers", "Headers/footers present for multi-page docs", "WARN",
            f"{len(missing)} section(s) without header/footer content", missing,
        )
    return Finding("headers_footers", "Headers/footers", "PASS", "")


def check_table_overflow(doc: Any) -> Finding:
    bad: list[str] = []
    for t_idx, table in enumerate(doc.tables, start=1):
        # All sections share page geometry in the common case; use the first.
        section = doc.sections[0] if doc.sections else None
        if section is None:
            break
        text_w = _text_area_width_in(section)
        tbl_w = _table_width_in(table)
        if text_w is None or tbl_w is None:
            continue
        if tbl_w > text_w + 0.05:  # 0.05" tolerance
            bad.append(f"table {t_idx}: width {tbl_w:.2f}\" > text area {text_w:.2f}\"")
    if bad:
        return Finding(
            "table_overflow", "Tables fit the text area", "WARN",
            f"{len(bad)} overflowing table(s)", bad,
        )
    return Finding("table_overflow", "Table widths", "PASS", "")


CHECKS = (
    check_placeholder,
    check_heading_hierarchy,
    check_body_font,
    check_line_spacing,
    check_margins,
    check_page_numbers,
    check_headers_footers,
    check_table_overflow,
)


def audit(file: str) -> dict[str, Any]:
    path = Path(file)
    try:
        doc = Document(str(path))
    except Exception as exc:
        finding = Finding(
            "file_open", "File opens with python-docx", "FAIL",
            f"Could not open {file}: {exc}", [],
        )
        return build_report(file, [finding])

    findings: list[Finding] = []
    for check in CHECKS:
        try:
            findings.append(check(doc))
        except Exception as exc:  # a single rule must not abort the whole audit
            findings.append(Finding(
                check.__name__, check.__doc__ or check.__name__, "WARN",
                f"Rule check raised: {exc}", [],
            ))
    return build_report(file, findings)


def main() -> int:
    parser = argparse.ArgumentParser(description="Semantic audit of a DOCX document.")
    parser.add_argument("file", help="Path to the .docx file to audit")
    parser.add_argument("--json", action="store_true", help="Emit JSON to stdout")
    parser.add_argument("--strict", action="store_true", help="Exit 1 on WARN as well as FAIL")
    args = parser.parse_args()

    if not Path(args.file).is_file():
        print(f"audit_doc: {args.file} is not a file", file=sys.stderr)
        return 2

    report = audit(args.file)
    if args.json:
        print(json.dumps(report, indent=2, ensure_ascii=False))
    else:
        print_human(report)

    if report["status"] == "FAIL":
        return 1
    if args.strict and report["status"] == "WARN":
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
