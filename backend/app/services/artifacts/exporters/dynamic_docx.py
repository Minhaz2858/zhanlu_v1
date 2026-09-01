"""Dynamic DOCX renderer — executes a :class:`DocumentPlan` with a premium
design system.

This is the antithesis of a fixed template: the renderer has NO hard-coded
section order.  It walks ``plan.blocks`` and dispatches each block by
``type`` to a small drawing routine.  The SAME engine renders an executive
one-pager or a 30-page analytical dossier — the difference is entirely in
the blocks the architect / LLM produced, not in this code.

Design language (modern AI-agent quality, à la Kimi / MiniMax):
  * Accent color band on the cover, section dividers with accent rules.
  * KPI "cards" with accent top-borders and large values.
  * Shaded callout boxes with a colored left rule (info / risk / opportunity).
  * Clean data tables: filled header, zebra rows, hairline horizontal rules.
  * Real charts (bar / line / pie / donut) rendered with matplotlib and
    embedded as images, themed to the palette.
  * Pull-quotes, comparison grids, timelines, recommendations as check-cards.

Entry point: ``render_document_plan(plan, ctx) -> (bytes, mime, ext)``.
"""
from __future__ import annotations

import io
import logging
import os
import re
from typing import Optional

from app.services.artifacts.document_plan import DocumentBlock, DocumentPlan
from app.services.artifacts.exporters._common import ExportContext
from app.services.artifacts.exporters._theme import resolve_ctx_theme

logger = logging.getLogger(__name__)

# Keep matplotlib's cache off the (often read-only) home dir when running
# inside a container as a non-owner user.  A writable tmp dir avoids the
# "mkdir failed for /home/.../.cache/matplotlib: Permission denied" warning
# and the per-render fallback churn.
if not os.environ.get("MPLCONFIGDIR") or not os.access(
    os.environ["MPLCONFIGDIR"], os.W_OK
):
    _mpl_cache = "/tmp/matplotlib_cache"
    try:
        os.makedirs(_mpl_cache, exist_ok=True)
        os.environ["MPLCONFIGDIR"] = _mpl_cache
    except OSError:
        pass

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXT = ".docx"

# Variant → accent color key on the DeckTheme hex dict.
_VARIANT_ACCENT = {
    "info": "primary",
    "success": "delta_up",
    "warning": "warn_accent",
    "risk": "delta_down",
    "opportunity": "delta_up",
}


# ---------------------------------------------------------------------------
# docx XML helpers
# ---------------------------------------------------------------------------

def _docx_rgb(hex_str: str):
    from docx.shared import RGBColor
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill.lstrip("#").upper())
    tcPr.append(shd)


def _set_cell_border(cell, edge: str, color_hex: str, sz: int = 12) -> None:
    """Add a single colored border edge to a table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = borders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}")
        borders.append(el)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color_hex.lstrip("#").upper())


def _set_para_left_border(p, color_hex: str, sz: int = 24) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = pBdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pBdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex.lstrip("#").upper())


def _add_field(paragraph, field_code: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    for typ, txt in (("begin", None), ("separate", None), ("end", None)):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), typ)
        run._r.append(fc)
        if typ == "separate":
            t = OxmlElement("w:t")
            t.text = "1"
            run._r.append(t)
    instr = paragraph._p.findall(".//" + qn("w:instrText"))
    if instr:
        instr[0].text = f" {field_code} "


# ---------------------------------------------------------------------------
# Main entry
# ---------------------------------------------------------------------------

def render_document_plan(
    plan: DocumentPlan,
    ctx: Optional[ExportContext] = None,
) -> tuple[bytes, str, str]:
    """Render a DocumentPlan into a .docx. Returns (bytes, mime, ext)."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ctx = ctx or ExportContext()
    theme = resolve_ctx_theme(ctx)
    hx = theme.as_hex_dict()
    accent = (plan.accent or hx["primary"]).lstrip("#")
    accent_hex = f"#{accent}"
    doc = Document()

    # Base typography
    normal = doc.styles["Normal"]
    normal.font.name = theme.font_body
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for lvl, (size, color) in {1: (20, hx["text"]), 2: (15, hx["primary_dark"]),
                               3: (12.5, hx["muted"])}.items():
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.size = Pt(size)
        hs.font.color.rgb = _docx_rgb(color)
        hs.font.bold = True
        hs.font.name = theme.font_heading

    counter = {"section": 0}
    for block in plan.blocks:
        _render_block(doc, block, theme, hx, accent_hex, counter)

    _add_footer(doc, hx["muted"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), MIME, EXT


# ---------------------------------------------------------------------------
# Block dispatcher
# ---------------------------------------------------------------------------

def _render_block(doc, b: DocumentBlock, theme, hx, accent_hex, counter):
    t = b.type
    if t == "cover":
        _doc_cover(doc, b, hx, accent_hex, theme)
    elif t == "section_divider":
        counter["section"] += 1
        _doc_section(doc, b, hx, accent_hex, counter["section"])
    elif t == "heading":
        doc.add_heading(b.title or "", level=b.level or 2)
    elif t == "paragraph":
        _doc_paragraph(doc, b, hx, accent_hex)
    elif t == "bullets":
        _doc_bullets(doc, b)
    elif t == "numbered":
        _doc_numbered(doc, b)
    elif t == "kpi_grid":
        _doc_kpi_grid(doc, b, hx)
    elif t == "data_table":
        _doc_data_table(doc, b, hx)
    elif t == "chart":
        _doc_chart(doc, b, theme, hx)
    elif t == "callout":
        _doc_callout(doc, b, hx)
    elif t == "comparison":
        _doc_comparison(doc, b, hx, accent_hex)
    elif t == "timeline":
        _doc_timeline(doc, b, hx, accent_hex)
    elif t == "quote":
        _doc_quote(doc, b, hx, accent_hex)
    elif t == "findings":
        _doc_findings(doc, b, hx)
    elif t == "recommendations":
        _doc_recommendations(doc, b, hx, accent_hex)
    elif t == "methodology":
        _doc_callout(doc, DocumentBlock(type="callout", variant="info",
                                         title=b.title or "Methodology",
                                         text=b.text), hx)
    elif t == "appendix":
        _doc_data_table(doc, b, hx, caption=b.title or "Appendix")
    elif t == "image" and b.image:
        _doc_image(doc, b)
    else:
        if b.title:
            doc.add_heading(b.title, level=2)
        if b.text:
            doc.add_paragraph(b.text)


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _doc_cover(doc, b, hx, accent_hex, theme):
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Accent band across the top of the page.
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = band.rows[0].cells[0]
    _set_cell_bg(cell, accent_hex)
    cell.paragraphs[0].text = ""
    tr = cell._tc.getparent()  # <w:tr>
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), "360")
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)
    _set_cell_border(cell, "bottom", accent_hex, 0)
    _set_cell_border(cell, "top", accent_hex, 0)
    _set_cell_border(cell, "left", accent_hex, 0)
    _set_cell_border(cell, "right", accent_hex, 0)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(b.title or "Report")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = _docx_rgb(hx["text"])
    r.font.name = theme.font_heading
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)

    # Thin accent rule under the title.
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("━" * 6)
    rr.font.color.rgb = _docx_rgb(accent_hex)
    rr.font.size = Pt(18)
    rule.paragraph_format.space_after = Pt(10)

    if b.subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(b.subtitle)
        sr.font.size = Pt(13)
        sr.font.color.rgb = _docx_rgb(hx["muted"])
        s.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Generated by Zhanlu AI")
    mr.font.size = Pt(10)
    mr.font.color.rgb = _docx_rgb(hx["muted"])
    mr.italic = True

    doc.add_page_break()


def _doc_section(doc, b, hx, accent_hex, n):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    kicker = doc.add_paragraph()
    kr = kicker.add_run(f"SECTION {n:02d}")
    kr.font.size = Pt(9)
    kr.font.bold = True
    kr.font.color.rgb = _docx_rgb(accent_hex)
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(2)

    h = doc.add_paragraph()
    hr = h.add_run(b.title or "")
    hr.font.size = Pt(18)
    hr.font.bold = True
    hr.font.color.rgb = _docx_rgb(hx["text"])
    hr.font.name = "Inter"
    h.paragraph_format.space_after = Pt(4)
    # Accent bottom rule.
    _set_para_left_border(h, accent_hex, 0)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = h._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), accent_hex.lstrip("#").upper())
    if b.subtitle:
        sub = doc.add_paragraph()
        sr = sub.add_run(b.subtitle)
        sr.font.size = Pt(11)
        sr.font.color.rgb = _docx_rgb(hx["muted"])
        sr.italic = True
        sub.paragraph_format.space_after = Pt(8)


def _doc_paragraph(doc, b, hx, accent_hex):
    from docx.shared import Pt
    title = b.title or b.style.get("title")
    if title:
        hh = doc.add_heading(title, level=b.level or 2)
    p = doc.add_paragraph()
    text = b.text or ""
    run = p.add_run(text)
    if b.style.get("lead"):
        run.font.size = Pt(12)
        run.font.color.rgb = _docx_rgb(hx["text"])
        _set_para_left_border(p, accent_hex, 18)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(10)


def _doc_bullets(doc, b):
    for item in b.bullets or []:
        doc.add_paragraph(str(item), style="List Bullet")


def _doc_numbered(doc, b):
    for item in b.bullets or []:
        doc.add_paragraph(str(item), style="List Number")


def _doc_kpi_grid(doc, b, hx):
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    items = b.items or []
    if not items:
        return
    n = min(len(items), 4)
    cols = n if n <= 4 else 2
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows * 2, cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    idx = 0
    for r in range(rows):
        for c in range(cols):
            if idx >= len(items):
                break
            item = items[idx]
            idx += 1
            label = str(item.get("label") or item.get("name") or "")
            value = str(item.get("value") if item.get("value") is not None else item.get("display") or "")
            delta = item.get("delta") or ""
            caption = item.get("caption") or ""
            # Label row
            lc = table.rows[r * 2].cells[c]
            lc.text = ""
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = lp.add_run(label.upper())
            lr.font.size = Pt(8.5)
            lr.font.bold = True
            lr.font.color.rgb = _docx_rgb(hx["muted"])
            # Value row
            vc = table.rows[r * 2 + 1].cells[c]
            vc.text = ""
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = vp.add_run(value)
            vr.font.size = Pt(22)
            vr.font.bold = True
            vr.font.color.rgb = _docx_rgb(hx["text"])
            if delta:
                dp = vc.add_paragraph()
                dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dr = dp.add_run(str(delta))
                dr.font.size = Pt(9)
                dr.font.bold = True
                dr.font.color.rgb = _docx_rgb(hx["delta_up"])
            if caption:
                cp = vc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(str(caption))
                cr.font.size = Pt(8)
                cr.font.italic = True
                cr.font.color.rgb = _docx_rgb(hx["muted"])
            # Card styling: surface fill + accent top border.
            _set_cell_bg(lc, hx["surface"])
            _set_cell_bg(vc, hx["surface"])
            _set_cell_border(lc, "top", hx["primary"], 18)
            _set_cell_border(lc, "bottom", hx["border"], 4)
            _set_cell_border(vc, "bottom", hx["border"], 4)
            _set_cell_border(lc, "left", hx["border"], 4)
            _set_cell_border(lc, "right", hx["border"], 4)
            _set_cell_border(vc, "left", hx["border"], 4)
            _set_cell_border(vc, "right", hx["border"], 4)
    doc.add_paragraph()


def _doc_data_table(doc, b, hx, caption=None):
    from docx.shared import Pt
    columns = b.columns or []
    rows = b.rows or []
    if not columns and rows and isinstance(rows[0], dict):
        columns = list(rows[0].keys())
    if not columns:
        return
    title = b.title or caption
    if title:
        doc.add_heading(title, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Table Grid"
    # Header
    for j, col in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        _set_cell_bg(cell, hx["primary_dark"])
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = _docx_rgb("#ffffff")
    # Body with zebra striping
    for i, row in enumerate(rows):
        if isinstance(row, dict):
            vals = [row.get(c, "") for c in columns]
        else:
            vals = list(row)
        for j in range(len(columns)):
            cell = table.rows[i + 1].cells[j]
            cell.text = "" if j >= len(vals) else str(vals[j])
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = _docx_rgb(hx["text"])
            if i % 2 == 1:
                _set_cell_bg(cell, hx["surface"])
    doc.add_paragraph()


def _doc_chart(doc, b, theme, hx):
    from docx.shared import Inches
    title = b.title
    if title:
        doc.add_heading(title, level=2)
    png = _render_chart_png(b, theme, hx)
    if png:
        doc.add_picture(png, width=Inches(6.2))
        doc.paragraphs[-1].alignment = 1  # center
        doc.add_paragraph()


def _doc_callout(doc, b, hx):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    variant = b.variant or "info"
    accent_key = _VARIANT_ACCENT.get(variant, "primary")
    accent_hex = hx.get(accent_key, hx["primary"])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, _tint(accent_hex, 0.90))
    _set_cell_border(cell, "left", accent_hex, 36)
    cell.paragraphs[0].text = ""
    if b.title:
        th = cell.paragraphs[0]
        tr = th.add_run(b.title)
        tr.font.bold = True
        tr.font.size = Pt(11)
        tr.font.color.rgb = _docx_rgb(accent_hex)
    body = cell.add_paragraph()
    br = body.add_run(b.text or "")
    br.font.size = Pt(10)
    br.font.color.rgb = _docx_rgb(hx["text"])
    doc.add_paragraph()


def _doc_comparison(doc, b, hx, accent_hex):
    from docx.shared import Pt
    title = b.title or "Comparison"
    doc.add_heading(title, level=2)
    items = b.items or []
    if not items:
        return
    # Two-column: detect a/b keys.
    sample = items[0]
    if "a" in sample and "b" in sample:
        headers = ["", str(sample.get("a_label", "A")), str(sample.get("b_label", "B"))]
        rows = [[i.get("label", ""), str(i.get("a", "")), str(i.get("b", ""))] for i in items]
        cols = 3
    else:
        headers = [str(sample.get("label", "Metric")), str(sample.get("value", "Value"))]
        rows = [[i.get("label", ""), str(i.get("value", ""))] for i in items]
        cols = 2
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    for j, hcol in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = hcol
        _set_cell_bg(cell, hx["primary_dark"])
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = _docx_rgb("#ffffff")
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(row[j])
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = _docx_rgb(hx["text"])
    doc.add_paragraph()


def _doc_timeline(doc, b, hx, accent_hex):
    from docx.shared import Pt
    doc.add_heading(b.title or "Timeline", level=2)
    for item in b.items or []:
        p = doc.add_paragraph()
        date = item.get("date", "")
        title = item.get("title", "")
        desc = item.get("desc", "")
        dr = p.add_run(f"{date}  ")
        dr.font.bold = True
        dr.font.color.rgb = _docx_rgb(accent_hex)
        dr.font.size = Pt(11)
        tr = p.add_run(title)
        tr.font.bold = True
        tr.font.size = Pt(11)
        if desc:
            dp = doc.add_paragraph()
            dp.add_run(desc).font.size = Pt(10)
            dp.paragraph_format.left_indent = Pt(14)


def _doc_quote(doc, b, hx, accent_hex):
    from docx.shared import Pt
    p = doc.add_paragraph()
    r = p.add_run(f"\u201c{b.text}\u201d")
    r.font.size = Pt(15)
    r.font.italic = True
    r.font.color.rgb = _docx_rgb(hx["text"])
    _set_para_left_border(p, accent_hex, 36)
    p.paragraph_format.left_indent = Pt(8)
    if b.title:
        cap = doc.add_paragraph()
        cr = cap.add_run(f"— {b.title}")
        cr.font.size = Pt(10)
        cr.font.color.rgb = _docx_rgb(hx["muted"])
        cr.italic = True


def _doc_findings(doc, b, hx):
    doc.add_heading(b.title or "Key Findings", level=2)
    for item in b.items or []:
        label = item.get("label") or ""
        text = item.get("text") or ""
        p = doc.add_paragraph(style="List Bullet")
        if label:
            lr = p.add_run(f"{label}: ")
            lr.font.bold = True
            lr.font.color.rgb = _docx_rgb(hx["finding_accent"])
        p.add_run(text)


def _doc_recommendations(doc, b, hx, accent_hex):
    from docx.shared import Pt
    doc.add_heading(b.title or "Recommendations", level=2)
    for item in b.items or []:
        label = item.get("label") or ""
        text = item.get("text") or ""
        p = doc.add_paragraph()
        cr = p.add_run("✔  ")
        cr.font.color.rgb = _docx_rgb(hx["delta_up"])
        cr.font.bold = True
        if label:
            lr = p.add_run(f"{label}: ")
            lr.font.bold = True
        p.add_run(text)


def _doc_image(doc, b):
    from docx.shared import Inches
    import base64
    src = b.image
    try:
        if src.startswith("data:"):
            _, b64 = src.split(",", 1)
            data = base64.b64decode(b64)
            doc.add_picture(io.BytesIO(data), width=Inches(6.0))
        else:
            doc.add_picture(src, width=Inches(6.0))
        doc.paragraphs[-1].alignment = 1
    except Exception as e:
        logger.warning("dynamic_docx: failed to embed image: %s", e)


# ---------------------------------------------------------------------------
# Chart rendering (matplotlib → PNG)
# ---------------------------------------------------------------------------

def _render_chart_png(b: DocumentBlock, theme, hx) -> Optional[bytes]:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as e:  # pragma: no cover
        logger.warning("dynamic_docx: matplotlib unavailable: %s", e)
        return None

    chart = b.chart or {}
    ctype = (b.chart_type or chart.get("type") or "bar").lower()
    palette = theme.chart_palette or [hx["primary"]]
    fig, ax = plt.subplots(figsize=(6.4, 3.2), dpi=200)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    try:
        if ctype in ("pie", "donut"):
            labels = list(chart.get("x") or [])
            y = list(chart.get("y") or [])
            wedges, _ = ax.pie(y, labels=None, colors=palette,
                               autopct=lambda p: f"{p:.0f}%", pctdistance=0.8,
                               startangle=90, wedgeprops={"width": 0.42 if ctype == "donut" else 1.0})
            ax.legend(wedges, labels, loc="center left", bbox_to_anchor=(1.0, 0.5),
                      fontsize=8, frameon=False)
            ax.set_title(b.title or "", fontsize=11, fontweight="bold")
        elif ctype == "line":
            x = list(chart.get("x") or [])
            y = chart.get("y") or []
            if isinstance(y, list) and y and isinstance(y[0], list):
                for i, series in enumerate(y):
                    ax.plot(x, series, marker="o", color=palette[i % len(palette)], linewidth=2)
            else:
                ax.plot(x, y, marker="o", color=palette[0], linewidth=2)
            ax.set_xlabel(chart.get("x_label", ""), fontsize=8)
            ax.set_ylabel(chart.get("y_label", ""), fontsize=8)
        else:  # bar / stacked_bar
            x = list(chart.get("x") or [])
            y = chart.get("y") or []
            ax.bar(x, y, color=palette[0], width=0.6)
            ax.set_xlabel(chart.get("x_label", ""), fontsize=8)
            ax.set_ylabel(chart.get("y_label", ""), fontsize=8)
            plt.setp(ax.get_xticklabels(), rotation=35, ha="right", fontsize=7)
        ax.tick_params(axis="both", labelsize=8)
        for s in ("top", "right"):
            ax.spines[s].set_visible(False)
        ax.grid(axis="y", color=hx["border"], linewidth=0.6, alpha=0.7)
    except Exception as e:
        logger.warning("dynamic_docx: chart draw failed: %s", e)
        plt.close(fig)
        return None

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _add_footer(doc, muted_hex):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    section = doc.sections[0]
    footer = section.footer
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Generated by Zhanlu AI · Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = _docx_rgb(muted_hex)
    _add_field(para, "PAGE")
    para.add_run(" of ")
    _add_field(para, "NUMPAGES")


# ---------------------------------------------------------------------------
# color util
# ---------------------------------------------------------------------------

def _tint(hex_base: str, frac: float) -> str:
    h = hex_base.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    r = int(r + (255 - r) * frac)
    g = int(g + (255 - g) * frac)
    b = int(b + (255 - b) * frac)
    return f"#{r:02x}{g:02x}{b:02x}"


__all__ = ["render_document_plan"]
