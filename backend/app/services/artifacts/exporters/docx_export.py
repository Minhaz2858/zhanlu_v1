"""DOCX exporter — render a ReportCardPayload into a .docx file.

Uses python-docx as the primary renderer, falling back to pandoc
(HTML → DOCX) when python-docx is unavailable. This matches the
approach described in the bundled ``skills/docx/SKILL.md``.

Public entry point: ``render(payload, ctx) -> (bytes, mime, ext)``

**Disk-write audit (2026-07-15):**  The pandoc fallback path (``_render_via_pandoc``)
writes HTML to a ``NamedTemporaryFile`` (``delete=False``) and converts to DOCX.
Both files are cleaned in a ``finally`` block.  The python-docx path is fully
in-memory (``BytesIO``).  No persistent disk writes remain.
"""

from __future__ import annotations

import io
import logging
import subprocess
import tempfile
import os
from typing import Optional

from app.services.synexia.contracts import ReportCardPayload
from app.services.artifacts.exporters._common import (
    ExportContext, stamp_filename,
    chart_rows, chart_x_key, chart_y_keys,
)
from app.services.artifacts.exporters._theme import resolve_ctx_theme

logger = logging.getLogger(__name__)

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXT = ".docx"


def _val(obj, *keys):
    """Safely read a value from a Pydantic model or plain dict.

    Tries each *key* in order: ``getattr`` for Pydantic models, ``.get()`` for dicts.
    Returns ``""`` if nothing matches.
    """
    for k in keys:
        if isinstance(obj, dict):
            if k in obj:
                return obj[k]
        else:
            try:
                return getattr(obj, k)
            except (AttributeError, TypeError):
                continue
    return ""


def _looks_like_markdown_report(summary: str) -> bool:
    """True when the summary is a full deterministic report (markdown with
    ``## `` headings and ``|`` tables) rather than a plain narrative. The
    frontend uses the same signal (MessageBubble.jsx suppresses the raw
    DataTableCard when the message contains "## " headers)."""
    if not summary:
        return False
    return ("## " in summary) and ("|" in summary)


def _add_markdown_report(doc, markdown: str) -> None:
    """Render a deterministic business-report markdown (## headings, pipe
    tables, **bold**, bullets) as proper Word elements instead of a literal
    markdown blob. Kept deliberately small — it covers the exact subset
    emitted by business_reports.py / report builders."""
    from docx.shared import Pt
    import re as _re

    lines = (markdown or "").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        # Heading 1 (rare — title)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(_re.sub(r"\*\*", "", stripped[2:].strip()), level=1)
            i += 1
            continue
        # Heading 2 (## Section)
        if stripped.startswith("## "):
            doc.add_heading(_re.sub(r"\*\*", "", stripped[3:].strip()), level=2)
            i += 1
            continue
        # Pipe table: header row + separator + body rows
        if stripped.startswith("|") and stripped.endswith("|"):
            header_cells = [
                c.strip() for c in stripped.strip("|").split("|")
            ]
            # separator row: |---|---| (all dashes)
            if i + 1 < len(lines) and _re.match(
                r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]
            ) and all(
                set(c.strip().replace(":", "").replace("-", "")) == set()
                for c in lines[i + 1].strip().strip("|").split("|")
            ):
                i += 1  # skip separator
                body = []
                while i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                    body.append(
                        [c.strip() for c in lines[i + 1].strip().strip("|").split("|")]
                    )
                    i += 1
                ncols = len(header_cells)
                t = doc.add_table(rows=1 + len(body), cols=ncols, style="Light Grid Accent 1")
                for j, hc in enumerate(header_cells):
                    cell = t.rows[0].cells[j]
                    cell.text = _re.sub(r"\*\*", "", hc)
                    _shade_cell(cell, "F1F5F9")
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)
                            r.font.bold = True
                for ri, brow in enumerate(body, 1):
                    for j in range(ncols):
                        val = brow[j] if j < len(brow) else ""
                        cell = t.rows[ri].cells[j]
                        cell.text = _re.sub(r"\*\*", "", val)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
                doc.add_paragraph()
                i += 1
                continue
            # Not a table (single pipe row) — fall through to paragraph
        # Bullet
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_re.sub(r"\*\*", "", stripped[2:].strip()))
            i += 1
            continue
        # Numbered
        if _re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(_re.sub(r"\*\*", "", stripped.split(". ", 1)[1]))
            i += 1
            continue
        # Plain paragraph (strip leftover markdown emphasis)
        text = _re.sub(r"\*\*", "", stripped)
        text = _re.sub(r"`([^`]*)`", r"\1", text)
        if text:
            p = doc.add_paragraph(text)
            for r in p.runs:
                r.font.size = Pt(10)
        i += 1


def _docx_rgb(theme_color) -> "object":
    """Convert a pptx ``RGBColor`` (from DeckTheme) to a docx ``RGBColor``."""
    from docx.shared import RGBColor

    return RGBColor(theme_color[0], theme_color[1], theme_color[2])


def _shade_cell(cell, hex_fill: str) -> None:
    """Apply a background fill to a table cell via OOXML shading."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill.lstrip("#").upper())
    tcPr.append(shd)


def _render_via_python_docx(payload: ReportCardPayload, ctx: Optional[ExportContext] = None) -> bytes:
    """Primary: build a Claude-style .docx with python-docx.

    The document is organized as:

    1. Cover page — title, subtitle (source + generated_at), brand
       (``doc_type="report"`` only; ``"memo"`` gets a To/From/Date/Subject
       header block instead, ``"brief"`` skips straight to the body)
    2. Executive Summary — payload.summary
    3. Methodology — payload.methodology (if present)
    4. Key Metrics — payload.kpis as a 1×N table with bold values
    5. Key Findings — payload.key_findings as one paragraph per finding
    6. Insights — payload.insights as a bullet list
    7. Recommendations — payload.recommendations as a bullet list
    8. Custom sections — payload.sections in order
    9. Data — payload.chart rows as a table
    10. SQL — payload.sql in a shaded code block
    11. Footer — "Generated by Zhanlu AI · Page X of Y"

    (payload.next_step is deliberately excluded: it is conversational
    guidance for the chat user, not report content.)

    Colors and fonts come from the resolved ``DeckTheme`` (Phase C) —
    the default ``zhanlu-blue`` theme reproduces the legacy palette.
    Proper Word styles ("Heading 1", "Heading 2", "Heading 3") are
    used throughout so mammoth's docx→html conversion preserves the
    structure as <h1>/<h2>/<h3> (rather than a wall of <p> tags).
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ctx = ctx or ExportContext()
    theme = resolve_ctx_theme(ctx)
    doc_type = (ctx.doc_type or "report").strip().lower()
    if doc_type not in ("report", "brief", "memo"):
        doc_type = "report"

    # Resolved theme colors as docx RGBColor (Phase C).
    C_TEXT = _docx_rgb(theme.text)
    C_PRIMARY = _docx_rgb(theme.primary)
    C_PRIMARY_DARK = _docx_rgb(theme.primary_dark)
    C_MUTED = _docx_rgb(theme.muted)
    C_DELTA_UP = _docx_rgb(theme.delta_up)
    SURFACE_HEX = theme.as_hex_dict()["surface"]

    doc = Document()

    # ---- Base styles ----
    style = doc.styles["Normal"]
    style.font.name = theme.font_body
    style.font.size = Pt(11)

    # Slightly tighten line spacing for business-report look
    pf = style.paragraph_format
    pf.space_after = Pt(6)

    # Make sure heading 1/2/3 have a colored, larger font so they
    # render distinctly in mammoth's <h1>/<h2>/<h3> output.
    for level, (size, color) in {
        1: (18, C_TEXT),
        2: (14, C_PRIMARY_DARK),
        3: (12, C_MUTED),
    }.items():
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = color
        h_style.font.bold = True
        h_style.font.name = theme.font_heading

    # ---- 1. Front matter (doc-type aware) ----
    title = payload.title or "Zhanlu Report"
    if doc_type == "report":
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover_title.add_run("📊  " + title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = C_TEXT
        cover_title.paragraph_format.space_before = Pt(120)
        cover_title.paragraph_format.space_after = Pt(12)

        if payload.source:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sub.add_run(f"Source: {payload.source}")
            sr.font.size = Pt(13)
            sr.font.color.rgb = C_MUTED
            sub.paragraph_format.space_after = Pt(2)

        gen_at = payload.generated_at or ""
        if gen_at:
            sub2 = doc.add_paragraph()
            sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr2 = sub2.add_run(f"Generated {gen_at}")
            sr2.font.size = Pt(11)
            sr2.font.color.rgb = C_MUTED

        # Brand footer of cover page
        brand = doc.add_paragraph()
        brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand.paragraph_format.space_before = Pt(180)
        br = brand.add_run("Zhanlu AI · Intelligent Data Reports")
        br.font.size = Pt(10)
        br.font.color.rgb = C_MUTED
        br.italic = True

        # Page break to start the body
        doc.add_page_break()
        # (Table of contents removed 2026-08-24 per user request — the Word TOC
        # field only showed a "right-click to update" placeholder, and users did
        # not want it in any doc format. Documents go cover → body directly.)

    elif doc_type == "memo":
        _add_memo_header(doc, payload, title, C_TEXT, C_MUTED)
    else:  # brief — no front matter; title as the first heading
        h = doc.add_heading(title, level=1)
        if payload.generated_at:
            meta = doc.add_paragraph()
            mr = meta.add_run(f"Generated {payload.generated_at}")
            mr.font.size = Pt(9)
            mr.font.color.rgb = C_MUTED

    # ---- 3. Executive Summary ----
    if payload.summary:
        # When the summary is already a full deterministic business report
        # (markdown with ## headings — e.g. business_reports.py output), render
        # it as proper Word headings/tables instead of a literal markdown blob,
        # and SKIP the generic Key Metrics + raw Data sections below (the report
        # already contains its own Key Figures / Top Customers / Execution Risk
        # tables). Mirrors the frontend's suppression of the raw DataTableCard
        # when the message contains "## " headers.
        if _looks_like_markdown_report(payload.summary):
            _add_markdown_report(doc, payload.summary)
            _summary_is_report = True
        else:
            _summary_is_report = False
            h = doc.add_heading("Executive Summary", level=1)
            para = doc.add_paragraph(payload.summary)
            # Slight tint via paragraph spacing (full-color callout boxes
            # in python-docx require cell shading which is fiddly; rely on
            # Heading 1 styling + paragraph indent for visual separation).
            para.paragraph_format.left_indent = Cm(0.4)
            para.paragraph_format.space_after = Pt(12)
    else:
        _summary_is_report = False

    # ---- 3. Methodology ----
    if payload.methodology:
        doc.add_heading("Methodology", level=1)
        m = doc.add_paragraph(payload.methodology)
        m.paragraph_format.left_indent = Cm(0.4)

    # ---- 5. Key Metrics (KPI tiles as a 1×N table) ----
    # Skipped when the summary already IS a full report with its own
    # Key Figures table (generic KPI tiles duplicate it).
    if payload.kpis and not _summary_is_report:
        doc.add_heading("Key Metrics", level=1)
        num_kpis = len(payload.kpis)
        table = doc.add_table(rows=2, cols=num_kpis, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row (labels) — tinted with the theme surface color.
        for i, kpi in enumerate(payload.kpis):
            cell = table.rows[0].cells[i]
            cell.text = str(_val(kpi, "label", "name") or f"KPI {i + 1}")
            _shade_cell(cell, SURFACE_HEX)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.bold = True
                    r.font.color.rgb = C_MUTED
        # Value row (values + optional delta + caption)
        for i, kpi in enumerate(payload.kpis):
            cell = table.rows[1].cells[i]
            cell.text = ""  # clear default
            value_str = str(_val(kpi, "value", "display"))
            if value_str:
                vp = cell.paragraphs[0]
                vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                vr = vp.add_run(value_str)
                vr.font.size = Pt(20)
                vr.font.bold = True
                vr.font.color.rgb = C_TEXT
            delta = _val(kpi, "delta")
            if delta:
                dp = cell.add_paragraph()
                dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dr = dp.add_run(str(delta))
                dr.font.size = Pt(9)
                dr.font.bold = True
                dr.font.color.rgb = C_DELTA_UP
            caption = _val(kpi, "caption")
            if caption:
                cp = cell.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(str(caption))
                cr.font.size = Pt(8)
                cr.font.color.rgb = C_MUTED
                cr.italic = True
        doc.add_paragraph()

    # ---- 5. Key Findings (narrative paragraphs) ----
    if payload.key_findings:
        doc.add_heading("Key Findings", level=1)
        for finding in payload.key_findings:
            doc.add_paragraph(_val(finding, "text") or "")

    # ---- 7. Insights (bullets) ----
    if payload.insights:
        doc.add_heading("Insights", level=1)
        for ins in payload.insights:
            doc.add_paragraph(_val(ins, "text") or "", style="List Bullet")

    # ---- 8. Recommendations (bullets) ----
    if payload.recommendations:
        doc.add_heading("Recommendations", level=1)
        for rec in payload.recommendations:
            doc.add_paragraph(_val(rec, "text") or "", style="List Bullet")

    # ---- 9. Custom sections ----
    # payload.sections may be a list of pydantic ``SectionSpec`` (as produced
    # by ``_payload_to_reportcard``) OR a list of plain dicts (legacy callers).
    # BOTH shapes must render — the previous code only handled dicts with a
    # ``heading``/``content`` key, so pydantic ``SectionSpec`` items (which
    # carry ``title``/``content``) were silently dropped, collapsing the
    # document to a cover + bare Executive Summary. Tolerate the common key
    # aliases the agent/LLM uses:
    #   {title, content} · {heading, body} · {heading, paragraphs|bullets} · {name, text}
    for sec in payload.sections or []:
        if isinstance(sec, dict):
            heading = sec.get("title") or sec.get("heading") or sec.get("name")
            content = sec.get("content") or sec.get("body") or sec.get("text")
            bullets = list(sec.get("bullets") or sec.get("paragraphs") or [])
            stype = sec.get("type", "narrative") or "narrative"
        else:
            heading = getattr(sec, "title", None) or getattr(sec, "heading", None)
            content = getattr(sec, "content", None) or getattr(sec, "body", None)
            bullets = list(getattr(sec, "bullets", None) or [])
            stype = getattr(sec, "type", "narrative") or "narrative"

        if heading:
            doc.add_heading(str(heading), level=2)
        if content:
            if isinstance(content, list):
                for line in content:
                    doc.add_paragraph(str(line))
            else:
                doc.add_paragraph(str(content))
        for b in bullets:
            doc.add_paragraph(str(b), style="List Bullet")

    # ---- 10. Chart data table ----
    # Skipped when the summary already IS a full report (the report's own
    # Top Customers / Execution Risk tables carry the data — a raw 188-row
    # dump is noise, exactly what the frontend suppresses on "## " headers).
    chart_rows_data = chart_rows(payload) if not _summary_is_report else []
    if chart_rows_data:
        doc.add_heading(payload.chart.title or "Data", level=1)
        keys = list(chart_rows_data[0].keys()) if isinstance(chart_rows_data[0], dict) else []
        if not keys:
            keys = chart_y_keys(payload) + [chart_x_key(payload)]
        ncols = len(keys)
        # Capped to keep the table readable inside Word.
        nrows = min(len(chart_rows_data) + 1, 26)
        t = doc.add_table(rows=nrows, cols=ncols, style="Light Grid Accent 1")
        for j, key in enumerate(keys):
            tc = t.rows[0].cells[j]
            tc.text = str(key)
            _shade_cell(tc, SURFACE_HEX)
            for p in tc.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = C_MUTED
        for i, row in enumerate(chart_rows_data[:25], 1):
            if isinstance(row, dict):
                for j, key in enumerate(keys):
                    cell = t.rows[i].cells[j]
                    cell.text = str(row.get(key, ""))
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(10)
                            r.font.color.rgb = C_TEXT
            else:
                for j, val in enumerate(row[:ncols]):
                    cell = t.rows[i].cells[j]
                    cell.text = str(val)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(10)
                            r.font.color.rgb = C_TEXT
        doc.add_paragraph()

    # ---- 11. SQL block ----
    if payload.sql:
        doc.add_heading("SQL", level=1)
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(payload.sql)
        sql_run.font.name = "Consolas"
        sql_run.font.size = Pt(9)
        sql_run.font.color.rgb = C_PRIMARY_DARK
        # Light shading on the paragraph via XML — produces a tinted
        # code-block look without using tables.
        pPr = sql_para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), SURFACE_HEX.lstrip("#").upper())
        pPr.append(shd)

    # NOTE: payload.next_step is intentionally NOT rendered. It is
    # conversational guidance for the in-chat card (e.g. "try narrowing the
    # scope"), not report content — a report recipient should never read
    # instructions addressed to the chat user.

    _add_page_number_footer(doc, C_MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_memo_header(doc, payload: ReportCardPayload, title: str,
                     c_text, c_muted) -> None:
    """Memo doc_type front matter: To / From / Date / Subject block."""
    from docx.shared import Pt

    head = doc.add_paragraph()
    hr = head.add_run("MEMORANDUM")
    hr.font.size = Pt(16)
    hr.font.bold = True
    hr.font.color.rgb = c_text
    head.paragraph_format.space_after = Pt(12)

    rows = [
        ("To:", payload.source or "—"),
        ("From:", "Zhanlu AI"),
        ("Date:", payload.generated_at or "—"),
        ("Subject:", title),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label} ")
        lr.font.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = c_muted
        vr = p.add_run(str(value))
        vr.font.size = Pt(11)
        vr.font.color.rgb = c_text

    # Horizontal rule under the memo block (bottom border paragraph).
    rule = doc.add_paragraph()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_number_footer(doc, color=None) -> None:
    """Add a "Generated by Zhanlu AI · Page X of Y" footer to every page.

    Uses Word's PAGE and NUMPAGES fields so the count is correct in
    both Microsoft Word and mammoth's conversion.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if color is None:
        color = RGBColor(0x94, 0xA3, 0xB8)

    section = doc.sections[0]
    footer = section.footer
    # Clear default empty paragraph if present
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Generated by Zhanlu AI · Page "
    run = para.add_run("Generated by Zhanlu AI · Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = color

    # PAGE field
    _add_field(para, "PAGE")
    para.add_run(" of ").font.size = Pt(9)
    _add_field(para, "NUMPAGES")
    for r in para.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = color


def _add_field(paragraph, field_code: str) -> None:
    """Append a Word field (e.g. 'PAGE', 'NUMPAGES') to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_code} "
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    # Placeholder text shown if field isn't computed
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)


def _render_via_pandoc(payload: ReportCardPayload, ctx: Optional[ExportContext] = None) -> bytes:
    """Fallback: build HTML in-memory, convert to DOCX via pandoc."""
    # Build a simple HTML document (themed from ctx when provided)
    html = _build_html(payload, ctx)

    with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as hf:
        hf.write(html)
        html_path = hf.name

    docx_path = html_path.replace(".html", ".docx")
    try:
        subprocess.run(
            ["pandoc", "-f", "html", "-t", "docx", "--standalone", "-o", docx_path, html_path],
            check=True,
            capture_output=True,
            timeout=30,
        )
        with open(docx_path, "rb") as df:
            return df.read()
    finally:
        for p in (html_path, docx_path):
            try:
                os.unlink(p)
            except OSError:
                pass


def _build_html(payload: ReportCardPayload, ctx: Optional[ExportContext] = None) -> str:
    """Build a themed Claude-style HTML report from the payload data.

    Used as the input for the pandoc HTML→DOCX fallback path.  Kept
    structurally compatible with the python-docx output so the
    rendered document is the same regardless of the code path.
    """
    ctx = ctx or ExportContext()
    theme = resolve_ctx_theme(ctx)
    hexes = theme.as_hex_dict()
    parts = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head><meta charset='utf-8'><title>",
        payload.title or "Zhanlu Report",
        "</title><style>",
        f"body {{ font-family: {theme.font_body}, sans-serif; max-width: 800px; margin: 0 auto; padding: 2em; color: {hexes['text']}; }}",
        f"h1 {{ color: {hexes['text']}; border-bottom: 2px solid {hexes['primary']}; padding-bottom: 4px; }}",
        f"h1, h2, h3 {{ font-family: {theme.font_heading}, sans-serif; }}",
        f"h2 {{ color: {hexes['primary']}; }}",
        f"h3 {{ color: {hexes['muted']}; }}",
        "table { width: 100%; border-collapse: collapse; margin: 1em 0; }",
        f"th {{ background: {hexes['primary_dark']}; color: #fff; padding: 8px 12px; text-align: left; font-weight: 600; }}",
        f"td {{ padding: 6px 12px; border-bottom: 1px solid {hexes['border']}; }}",
        f"tr:nth-child(even) td {{ background: {hexes['surface']}; }}",
        ".kpi-table { width: 100%; border-collapse: collapse; margin: 1em 0; }",
        f".kpi-table td {{ border: 1px solid {hexes['border']}; padding: 12px; text-align: center; vertical-align: top; }}",
        f".kpi-table .label {{ font-size: 10px; color: {hexes['muted']}; text-transform: uppercase; letter-spacing: 0.5px; }}",
        f".kpi-table .value {{ font-size: 20px; font-weight: 700; color: {hexes['text']}; margin-top: 4px; }}",
        f".kpi-table .delta {{ font-size: 10px; font-weight: 600; color: {hexes['delta_up']}; background: {hexes['warn_bg']}; padding: 1px 5px; border-radius: 3px; display: inline-block; margin-top: 4px; }}",
        f".kpi-table .caption {{ font-size: 9px; color: {hexes['muted']}; margin-top: 2px; font-style: italic; }}",
        f".executive-summary {{ background: {hexes['insight_bg']}; border-left: 4px solid {hexes['finding_accent']}; padding: 12px 16px; margin: 1em 0; }}",
        f".methodology {{ background: {hexes['surface']}; border-left: 4px solid {hexes['muted']}; padding: 12px 16px; margin: 1em 0; }}",
        f".next-step {{ background: {hexes['next_bg']}; border-left: 4px solid {hexes['primary']}; padding: 12px 16px; margin: 1em 0; font-style: italic; }}",
        ".cover { text-align: center; padding: 60px 0; }",
        f".cover h1 {{ font-size: 32px; color: {hexes['text']}; border: none; margin: 0; }}",
        f".cover .source {{ color: {hexes['muted']}; font-size: 14px; margin-top: 12px; }}",
        f".cover .generated {{ color: {hexes['muted']}; font-size: 12px; margin-top: 4px; }}",
        f".cover .brand {{ color: {hexes['muted']}; font-size: 11px; margin-top: 80px; font-style: italic; }}",
        "ul { padding-left: 1.5em; }",
        "li { margin: 4px 0; line-height: 1.5; }",
        f"pre {{ background: {hexes['surface']}; padding: 10px; border-radius: 4px; font-family: Consolas, monospace; font-size: 11px; overflow-x: auto; color: {hexes['primary_dark']}; }}",
        f".footer {{ text-align: center; color: {hexes['muted']}; font-size: 11px; margin-top: 2em; }}",
        "</style></head><body>",
    ]

    # ---- Cover page ----
    parts.append("<div class='cover'>")
    parts.append(f"<h1>📊 {payload.title or 'Zhanlu Report'}</h1>")
    if payload.source:
        parts.append(f"<div class='source'>Source: {payload.source}</div>")
    if payload.generated_at:
        parts.append(f"<div class='generated'>Generated {payload.generated_at}</div>")
    parts.append("<div class='brand'>Zhanlu AI · Intelligent Data Reports</div>")
    parts.append("</div>")
    parts.append("<div style='page-break-after: always;'></div>")

    # ---- Executive Summary ----
    if payload.summary:
        parts.append("<h1>Executive Summary</h1>")
        parts.append(f"<div class='executive-summary'>{payload.summary}</div>")

    # ---- Methodology ----
    if payload.methodology:
        parts.append("<h1>Methodology</h1>")
        parts.append(f"<div class='methodology'>{payload.methodology}</div>")

    # ---- KPI tiles (1×N table) ----
    if payload.kpis:
        parts.append("<h1>Key Metrics</h1>")
        parts.append("<table class='kpi-table'><tr>")
        for kpi in payload.kpis:
            label = _val(kpi, "label", "name")
            value = _val(kpi, "value", "display")
            delta = _val(kpi, "delta")
            caption = _val(kpi, "caption")
            cell = "<td>"
            cell += f"<div class='label'>{label}</div>"
            cell += f"<div class='value'>{value}</div>"
            if delta:
                cell += f"<div class='delta'>{delta}</div>"
            if caption:
                cell += f"<div class='caption'>{caption}</div>"
            cell += "</td>"
            parts.append(cell)
        parts.append("</tr></table>")

    # ---- Key Findings (narrative paragraphs) ----
    if payload.key_findings:
        parts.append("<h1>Key Findings</h1>")
        for f in payload.key_findings:
            text = f if isinstance(f, str) else _val(f, "text", "label") or str(f)
            parts.append(f"<p>{text}</p>")

    # ---- Insights (bullets) ----
    if payload.insights:
        parts.append("<h1>Insights</h1><ul>")
        for insight in payload.insights:
            text = insight if isinstance(insight, str) else _val(insight, "text", "label") or str(insight)
            parts.append(f"<li>{text}</li>")
        parts.append("</ul>")

    # ---- Recommendations (bullets) ----
    if payload.recommendations:
        parts.append("<h1>Recommendations</h1><ul>")
        for rec in payload.recommendations:
            text = rec if isinstance(rec, str) else _val(rec, "text", "label") or str(rec)
            parts.append(f"<li>{text}</li>")
        parts.append("</ul>")

    # ---- Custom sections ----
    for sec in payload.sections:
        if not sec.title:
            continue
        parts.append(f"<h1>{sec.title}</h1>")
        if sec.content:
            parts.append(f"<p>{sec.content}</p>")
        for b in (sec.bullets or []):
            parts.append(f"<li>{b}</li>")

    # ---- Chart / Table ----
    if payload.chart and payload.chart.data:
        chart = payload.chart
        chart_title = _val(chart, "title", "label") or "Data"
        parts.append(f"<h1>{chart_title}</h1>")
        data_points = list(chart.data)
        if isinstance(data_points[0], dict):
            sample = data_points[0]
            keys = [k for k in sample.keys() if k.lower() in ("label", "name", "x", "category")] + \
                   [k for k in sample.keys() if k.lower() in ("value", "y", "amount", "count")]
            keys = list(dict.fromkeys(keys))
            if not keys:
                keys = list(sample.keys())[:4]
            parts.append("<table><thead><tr>")
            for key in keys:
                parts.append(f"<th>{key}</th>")
            parts.append("</tr></thead><tbody>")
            for pt in data_points:
                parts.append("<tr>")
                for key in keys:
                    parts.append(f"<td>{pt.get(key, '')}</td>")
                parts.append("</tr>")
            parts.append("</tbody></table>")

    # ---- SQL block ----
    if payload.sql:
        parts.append("<h1>SQL</h1>")
        parts.append(f"<pre>{payload.sql}</pre>")

    # next_step is intentionally omitted — see render_docx note above.

    parts.append('<p class="footer">Generated by Zhanlu AI</p>')
    parts.append("</body></html>")
    return "\n".join(parts)


def _extract_doc_plan(payload):
    """Return a DocumentPlan to render dynamically, or None for the legacy path."""
    from app.services.artifacts.document_plan import DocumentPlan

    try:
        if isinstance(payload, DocumentPlan):
            return payload
        if isinstance(payload, dict):
            if payload.get("blocks"):
                meta = {k: payload[k] for k in ("title", "subtitle", "theme") if payload.get(k)}
                return DocumentPlan.from_blocks(payload["blocks"], meta=meta)
            return None
        raw = getattr(payload, "blocks", None)
        if raw:
            meta = {
                "title": getattr(payload, "title", ""),
                "theme": getattr(payload, "theme", ""),
            }
            return DocumentPlan.from_blocks(list(raw), meta=meta)
    except Exception as exc:  # pragma: no cover
        logger.warning("docx: could not parse dynamic plan: %s", exc)
    return None


def render(
    payload: ReportCardPayload,
    ctx: Optional[ExportContext] = None,
) -> tuple[bytes, str, str]:
    """Render a ReportCardPayload into a .docx file.

    Returns:
        (bytes, mime_type, file_extension)

    Fully-dynamic path: when the payload carries an explicit ``blocks`` list
    (authored by the agent or produced by the server-side architect), the
    document is rendered by executing that ordered block plan — there is no
    fixed template.  Otherwise the legacy fixed layout (and the enterprise
    pipeline) apply.

    Phase 1C: When the payload carries ``enterprise_report_kind ==
    "executive"`` (the marker set by the ``collect_enterprise_data``
    tool), delegate to ``render_enterprise_docx`` to produce the full
    6-section executive document (cover, exec summary with citation
    anchors, KPI grid, segment breakdown, drivers, risks, actions,
    lineage appendix) instead of the generic ReportCard layout.
    """
    # --- Fully-dynamic block rendering (Phase: dynamic document generation) ---
    _plan = _extract_doc_plan(payload)
    if _plan is not None:
        try:
            from app.services.artifacts.exporters.dynamic_docx import render_document_plan
            # render_document_plan already returns the (bytes, mime, ext)
            # tuple — do NOT re-wrap it, or the caller gets a nested tuple
            # whose first element is itself a tuple instead of raw bytes.
            data, _mime, _ext = render_document_plan(_plan, ctx)
            logger.info(
                "DOCX rendered via dynamic plan (%d bytes, %d blocks)",
                len(data), len(_plan.blocks),
            )
            return data, MIME, EXT
        except Exception as e:
            logger.warning("dynamic DOCX render failed (%s); falling back", e)

    # Phase 1C: Enterprise pipeline short-circuit. Accept both
    # pydantic models and plain dicts (the orchestrator passes a
    # dict from the tool result).
    _is_enterprise = (
        (isinstance(payload, dict) and payload.get("enterprise_report_kind") == "executive")
        or (
            hasattr(payload, "enterprise_report_kind")
            and getattr(payload, "enterprise_report_kind", None) == "executive"
        )
    )
    if _is_enterprise:
        try:
            from app.services.enterprise_orchestrator.renderers import (
                render_enterprise_docx,
            )
            data = render_enterprise_docx(dict(payload) if not isinstance(payload, dict) else payload)
            logger.info(
                "DOCX rendered via enterprise pipeline (%d bytes, kind=%s)",
                len(data), (payload.get("enterprise_report_kind") if isinstance(payload, dict) else getattr(payload, "enterprise_report_kind", "?")),
            )
            return data, MIME, EXT
        except ImportError:
            logger.warning("enterprise_orchestrator renderer unavailable, falling back to generic DOCX")
        except Exception as _exc:
            logger.warning(
                "enterprise DOCX render failed (%s), falling back to generic",
                _exc,
            )

    # Try python-docx first (richer output).
    try:
        data = _render_via_python_docx(payload, ctx)
        logger.info("DOCX rendered via python-docx (%d bytes)", len(data))
        return data, MIME, EXT
    except ImportError:
        logger.info("python-docx not available, falling back to pandoc")
    except Exception as e:
        logger.warning("python-docx render failed (%s), falling back to pandoc", e)

    # Fallback: pandoc HTML → DOCX (theme is applied via inline style on body)
    try:
        data = _render_via_pandoc(payload, ctx)
        logger.info("DOCX rendered via pandoc (%d bytes)", len(data))
        return data, MIME, EXT
    except FileNotFoundError:
        logger.error("pandoc not found — neither python-docx nor pandoc is available")
        raise
    except Exception as e:
        logger.error("pandoc render failed: %s", e)
        raise
