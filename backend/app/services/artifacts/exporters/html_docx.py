"""HTML → DOCX renderer — converts standalone HTML bytes to a .docx file.

Two-stage strategy:
1. **pandoc** (preferred) — pipe HTML in, get .docx out.  Handles CSS,
   tables, images far better than python-docx.
2. **python-docx + BeautifulSoup** (fallback) — parse HTML structure and
   rebuild it with python-docx styles.  Best-effort; complex layouts will
   be simplified.

Public entry point: ``render_html_to_docx(html_bytes) -> bytes``

**Disk-write audit (2026-07-15):**  The pandoc path writes HTML and DOCX
to ``NamedTemporaryFile`` (``delete=False``), cleaned in a ``finally`` block.
The python-docx+bs4 path is fully in-memory (``BytesIO``).  No persistent
disk writes remain.
"""

from __future__ import annotations

import io
import logging
import os
import subprocess
import tempfile

logger = logging.getLogger(__name__)

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
EXT = ".docx"


def render_html_to_docx(html_bytes: bytes) -> bytes:
    """Convert HTML content to a .docx file.

    Tries pandoc first; falls back to python-docx + BeautifulSoup if
    pandoc is unavailable or fails.
    """
    # Try pandoc first — far richer output
    try:
        return _via_pandoc(html_bytes)
    except (FileNotFoundError, subprocess.CalledProcessError, OSError) as exc:
        logger.info("pandoc HTML→DOCX failed (%s), falling back to python-docx", exc)

    # Fallback: python-docx + BeautifulSoup
    try:
        return _via_python_docx(html_bytes)
    except ImportError as exc:
        logger.error("Neither pandoc nor python-docx+bs4 is available: %s", exc)
        raise RuntimeError("Cannot render HTML to DOCX — missing dependencies") from exc
    except Exception as exc:
        logger.error("python-docx HTML→DOCX fallback failed: %s", exc)
        raise


# ── Pandoc path ────────────────────────────────────────────────────────────


def _via_pandoc(html_bytes: bytes) -> bytes:
    """Pipe HTML bytes through pandoc to produce .docx."""
    # Write HTML to a temp file (pandoc reads from stdin but .docx needs a file)
    with tempfile.NamedTemporaryFile(
        mode="wb", suffix=".html", delete=False
    ) as hf:
        hf.write(html_bytes)
        html_path = hf.name

    docx_path = html_path.replace(".html", ".docx")
    try:
        subprocess.run(
            [
                "pandoc",
                "-f", "html",
                "-t", "docx",
                "--standalone",
                "-o", docx_path,
                html_path,
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )
        with open(docx_path, "rb") as df:
            result = df.read()
        logger.info("HTML→DOCX via pandoc: %d bytes", len(result))
        return result
    finally:
        for p in (html_path, docx_path):
            try:
                os.unlink(p)
            except OSError:
                pass


# ── python-docx + BeautifulSoup fallback ───────────────────────────────────


def _via_python_docx(html_bytes: bytes) -> bytes:
    """Build a .docx by parsing HTML structure with BeautifulSoup and
    translating semantic elements (h1–h6, p, ul, ol, table, pre, blockquote)
    into python-docx styles.

    Inline formatting is preserved: ``<strong>/<b>``, ``<em>/<i>`` and
    ``<code>`` become bold/italic/monospace runs instead of being flattened.
    CJK text renders in a CJK-capable font via the ``w:eastAsia`` attribute.
    """
    from bs4 import BeautifulSoup, NavigableString
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    CJK_FONT = "Microsoft YaHei"  # falls back gracefully on non-CJK systems

    soup = BeautifulSoup(html_bytes, "html.parser")
    doc = Document()

    # Normal style — Latin + CJK font so Chinese text renders correctly.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    # Mapping: tag → heading level
    HEADING_MAP = {
        "h1": 0, "h2": 1, "h3": 2, "h4": 3, "h5": 4, "h6": 5,
    }

    def _apply_font(run, *, name: str | None = None, cjk: bool = True):
        if name:
            run.font.name = name
        if cjk:
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), CJK_FONT)

    def _add_inline_runs(paragraph, el):
        """Walk an element's descendants and add runs preserving inline tags."""
        for node in el.descendants:
            if isinstance(node, NavigableString):
                text = str(node)
                if not text.strip():
                    # Preserve a single space so words don't merge; skip other ws.
                    if text:
                        run = paragraph.add_run(" ")
                        _apply_font(run)
                    continue
                # Determine formatting from ancestor tags.
                bold = italic = code = False
                for anc in node.parents:
                    name = (anc.name or "").lower() if hasattr(anc, "name") else ""
                    if name in ("strong", "b"):
                        bold = True
                    elif name in ("em", "i"):
                        italic = True
                    elif name == "code":
                        code = True
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    _apply_font(run, name="Courier New", cjk=False)
                    run.font.size = Pt(9)
                else:
                    _apply_font(run)

    def _add_paragraph(el, style_name: str | None = None):
        """Add a paragraph, preserving inline runs."""
        p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        _add_inline_runs(p, el)
        if not p.runs:
            p.add_run("")

    def _process_element(el):
        tag = (el.name or "").lower() if hasattr(el, "name") else ""

        if tag in HEADING_MAP:
            level = HEADING_MAP[tag]
            text = el.get_text(strip=True)
            if text:
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    _apply_font(run)

        elif tag == "p":
            if el.get_text(strip=True):
                _add_paragraph(el)

        elif tag in ("ul", "ol"):
            for li in el.find_all("li", recursive=False):
                if li.get_text(strip=True):
                    _add_paragraph(li, "List Bullet" if tag == "ul" else "List Number")

        elif tag == "table":
            rows = el.find_all("tr")
            if not rows:
                return
            # Determine column count from the widest row.
            cols = 0
            for row in rows:
                cols = max(cols, len(row.find_all(["th", "td"])))
            cols = max(cols, 1)
            table = doc.add_table(rows=len(rows), cols=cols, style="Light Grid Accent 1")
            for i, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for j, cell in enumerate(cells):
                    if j >= cols:
                        break
                    target = table.rows[i].cells[j]
                    # Clear the default empty paragraph, then add inline runs.
                    target.paragraphs[0].clear() if hasattr(target.paragraphs[0], "clear") else None
                    p = target.paragraphs[0]
                    _add_inline_runs(p, cell)
                    # Bold the header row.
                    if i == 0:
                        for run in p.runs:
                            run.bold = True
                    for run in p.runs:
                        _apply_font(run)

        elif tag == "pre":
            text = el.get_text()
            if text.strip():
                p = doc.add_paragraph()
                run = p.add_run(text)
                _apply_font(run, name="Courier New", cjk=False)
                run.font.size = Pt(9)

        elif tag == "blockquote":
            if el.get_text(strip=True):
                p = doc.add_paragraph(style="Quote")
                _add_inline_runs(p, el)

        elif tag == "hr":
            doc.add_paragraph("_" * 60)

        else:
            # Recurse into children for unknown containers (body, div, section, etc.)
            if hasattr(el, "children"):
                for child in el.children:
                    if hasattr(child, "name"):
                        _process_element(child)

    # Walk the body or top-level elements
    body = soup.find("body")
    if body:
        for child in body.children:
            if hasattr(child, "name"):
                _process_element(child)
    else:
        for child in soup.children:
            if hasattr(child, "name"):
                _process_element(child)

    # Add a page-number footer ("Generated by Zhanlu AI · Page X of Y").
    _add_page_number_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    result = buf.getvalue()
    logger.info("HTML→DOCX via python-docx+bs4: %d bytes", len(result))
    return result


def _add_page_number_footer(doc) -> None:
    """Add a centered "Generated by Zhanlu AI · Page X of Y" footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    color = RGBColor(0x94, 0xA3, 0xB8)
    section = doc.sections[0]
    footer = section.footer
    for p in list(footer.paragraphs):
        p._p.getparent().remove(p._p)

    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Generated by Zhanlu AI · Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = color

    _add_field(para, "PAGE")
    para.add_run(" of ").font.size = Pt(9)
    _add_field(para, "NUMPAGES")
    for r in para.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = color


def _add_field(paragraph, field_code: str) -> None:
    """Append a Word field (e.g. 'PAGE', 'NUMPAGES') to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def render(payload, ctx=None):
    """Adapter: ``render(payload, ctx) -> (bytes, mime, ext)`` so this module
    can be registered as an ExportService renderer alongside existing ones.

    ``payload`` is expected to be raw HTML bytes (not a ReportCardPayload).
    The caller (ExportService) wraps this when ``canonical_format=="html"``.
    """
    if isinstance(payload, bytes):
        html = payload
    elif isinstance(payload, str):
        html = payload.encode("utf-8")
    else:
        raise TypeError(f"html_docx render expects bytes or str, got {type(payload)}")
    return render_html_to_docx(html), MIME, EXT
