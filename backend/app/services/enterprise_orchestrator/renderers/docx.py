"""DOCX renderer for the enterprise business-data payload (spec §10.1).

Produces a 6-section executive .docx:

1. Cover page (title + period + source label)
2. Executive Summary (with claim citations as superscript markers)
3. Primary Metric Breakdown (table + KPI)
4. Segment Decomposition (table + observations)
5. Operational Drivers (domain-adaptive heading + list)
6. Risk Section (domain-adaptive heading + list)
7. Recommended Actions (numbered, severity-tagged)
8. Appendix: Data Lineage (per-facet SQL + execution log + row count)
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt

# Section order and their payload keys / labels.
_SECTIONS = [
    ("executive_summary", "Executive Summary"),
    ("primary_metric_breakdown", "Primary Metric Breakdown"),
    ("segment_decomposition", "Segment Decomposition"),
    ("operational_drivers", None),   # label is domain-adaptive
    ("risk_section", None),          # label is domain-adaptive
    ("recommended_actions", "Recommended Actions"),
]


def _table_columns(rows: list[dict]) -> list[str]:
    """Infer column order from the first appearance of keys across rows."""
    cols: list[str] = []
    for row in rows:
        if isinstance(row, dict):
            for key in row:
                if key not in cols:
                    cols.append(key)
    return cols


def _add_rows_table(doc: Document, rows: list[dict]) -> None:
    """Add a python-docx table for a list of dict rows."""
    if not rows:
        doc.add_paragraph("(data unavailable)")
        return
    cols = _table_columns(rows)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = str(col)
    for row in rows:
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            cells[i].text = "" if row.get(col) is None else str(row.get(col))


def _add_cited_summary(doc: Document, summary: str, claims: list[dict]) -> None:
    """Add the executive summary with claim citations as superscripts."""
    p = doc.add_paragraph()
    p.add_run(summary.strip() if summary else "")
    for i, claim in enumerate(claims or [], start=1):
        run = p.add_run(f"[{i}]")
        run.font.superscript = True
        run.font.size = Pt(9)


def _section_heading(doc: Document, section: dict, fallback: str) -> str:
    label = (section or {}).get("label") or fallback
    doc.add_heading(label, level=1)
    return label


def _add_list(doc: Document, items: list[str]) -> None:
    if not items:
        doc.add_paragraph("(data unavailable)")
        return
    for item in items:
        doc.add_paragraph(f"• {item}", style="List Bullet")


def _add_lineage(doc: Document, payload: dict) -> None:
    """Render the Data Lineage appendix from `lineage` (or `claims` fallback)."""
    doc.add_heading("Data Lineage", level=1)

    confidence = payload.get("data_confidence") or {}
    missing_reasons = confidence.get("missing_reasons") or {}
    missing_facets = confidence.get("missing_facets") or []

    lineage = payload.get("lineage")
    if lineage:
        for facet_id, info in lineage.items():
            doc.add_heading(f"Facet: {facet_id}", level=2)
            if not info.get("available"):
                reason = missing_reasons.get(facet_id, "data unavailable")
                doc.add_paragraph(f"(data unavailable: {reason})")
                continue
            if info.get("source_label"):
                doc.add_paragraph(f"Source: {info['source_label']}")
            if info.get("source_sql"):
                run = doc.add_paragraph().add_run(str(info["source_sql"]))
                run.font.name = "Courier New"
            doc.add_paragraph(f"Row count: {info.get('row_count', 0)}")
            for entry in info.get("execution_log") or []:
                step = entry.get("step", "?")
                latency = entry.get("latency_ms", "")
                status = entry.get("status", "")
                doc.add_paragraph(f"  • {step} ({latency} ms, {status})")
    else:
        claims = payload.get("claims") or []
        for i, claim in enumerate(claims, start=1):
            doc.add_heading(f"Claim {i}: {claim.get('claim_id', '?')}", level=2)
            if claim.get("source_sql"):
                run = doc.add_paragraph().add_run(str(claim["source_sql"]))
                run.font.name = "Courier New"

    # Surface any missing facets that were not already listed in `lineage`.
    listed = set((payload.get("lineage") or {}).keys())
    for facet_id in missing_facets:
        if facet_id in listed:
            continue
        doc.add_heading(f"Facet: {facet_id}", level=2)
        doc.add_paragraph(f"(data unavailable: {missing_reasons.get(facet_id, 'no data')})")


def render_enterprise_docx(payload: dict) -> bytes:
    """Render an enterprise payload into a .docx byte string."""
    doc = Document()

    # --- Cover page ---
    title = payload.get("title") or "Enterprise Data Report"
    doc.add_heading(title, level=0)
    if payload.get("period_label"):
        doc.add_paragraph(payload["period_label"])
    if payload.get("source_label"):
        doc.add_paragraph(f"Source: {payload['source_label']}")
    doc.add_page_break()

    # --- Section 1: Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    _add_cited_summary(doc, payload.get("executive_summary", ""), payload.get("claims") or [])

    # --- Section 2: Primary Metric Breakdown ---
    pmb = payload.get("primary_metric_breakdown") or {}
    _section_heading(doc, pmb, "Primary Metric Breakdown")
    if pmb.get("available"):
        kpi = pmb.get("kpi") or {}
        if kpi:
            kpi_text = "; ".join(f"{k}: {v}" for k, v in kpi.items())
            doc.add_paragraph(f"KPI: {kpi_text}")
        _add_rows_table(doc, pmb.get("rows") or [])
    else:
        doc.add_paragraph(f"(data unavailable: {pmb.get('unavailable_reason', '')})")

    # --- Section 3: Segment Decomposition ---
    seg = payload.get("segment_decomposition") or {}
    _section_heading(doc, seg, "Segment Decomposition")
    if seg.get("available"):
        _add_rows_table(doc, seg.get("rows") or [])
        for obs in seg.get("observations") or []:
            doc.add_paragraph(f"• {obs}", style="List Bullet")
    else:
        doc.add_paragraph(f"(data unavailable: {seg.get('unavailable_reason', '')})")

    # --- Section 4: Operational Drivers ---
    drivers = payload.get("operational_drivers") or {}
    _section_heading(doc, drivers, "Operational Drivers")
    if drivers.get("available"):
        _add_list(doc, drivers.get("drivers") or [])
    else:
        doc.add_paragraph("(data unavailable)")

    # --- Section 5: Risk Section ---
    risk = payload.get("risk_section") or {}
    _section_heading(doc, risk, "Risk Section")
    if risk.get("available"):
        _add_list(doc, risk.get("risks") or [])
    else:
        doc.add_paragraph("(data unavailable)")

    # --- Section 6: Recommended Actions ---
    doc.add_heading("Recommended Actions", level=1)
    actions = payload.get("recommended_actions") or []
    if not actions:
        doc.add_paragraph("(none)")
    for i, action in enumerate(actions, start=1):
        severity = action.get("severity", "medium")
        doc.add_paragraph(f"{i}. [{severity.upper()}] {action.get('action', '')}")

    # --- Appendix: Data Lineage ---
    _add_lineage(doc, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
